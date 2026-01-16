# --- Flask Pricing Calculator Application ---
# This app provides a pricing calculator for messaging services with dynamic inclusions, platform fees, and analytics.
# Key features: dynamic inclusions, robust error handling, session management, and professional UI.

from flask import Flask, render_template, request, session, redirect, url_for, flash, send_file
from calculator import calculate_pricing, get_suggested_price, meta_costs_table, calculate_total_mandays, calculate_total_manday_cost, COUNTRY_MANDAY_RATES, calculate_total_mandays_breakdown, get_committed_amount_rate_for_volume, get_lowest_tier_price
import os
import sys
import re
from collections import Counter, defaultdict
import statistics
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from datetime import datetime
from sqlalchemy import func
import uuid
from calculator import get_committed_amount_rates
from io import BytesIO
from docx import Document
from pricing_config import (
    committed_amount_slabs,
    PLATFORM_PRICING_GUIDANCE,
    VOICE_NOTES_PRICING,
    get_voice_notes_price,
    AI_AGENT_PRICING,
    compute_ai_price_components,
    get_default_location_for_email,
)
from calculator import get_committed_amount_rate_for_volume

# 🍕 PIZZA EASTER EGG SYSTEM 🍕
def generate_pizza_easter_egg_id(calculation_data=None):
    """
    Generate a fun pizza-themed calculation ID with achievement and time-based triggers!
    
    Args:
        calculation_data: Dict containing calculation details for achievement triggers
        
    Returns:
        str: Fun pizza-themed calculation ID
    """
    import random
    
    # 🍕 Pizza toppings for different calculation types
    pizza_toppings = {
        'basic': ['CHEESE', 'MARGHERITA', 'PLAIN-JANE', 'PANEER', 'TANDOORI'],
        'premium': ['PEPPERONI', 'SUPREME', 'DELUXE', 'CHICKEN-TIKKA', 'BUTTER-CHICKEN'],
        'luxury': ['TRUFFLE', 'CAVIAR', 'GOLD-LEAF', 'SAFFRON', 'GOLDEN-TEMPLE'],
        'funny': ['PINEAPPLE', 'ANCHOVY', 'BROCCOLI', 'PICKLE', 'CHUTNEY'],
        'achievement': ['WINNER', 'CHAMPION', 'HERO', 'MAHARAJA', 'SARDAR'],
        'time': ['BREAKFAST', 'LUNCH', 'DINNER', 'MIDNIGHT-SNACK', 'CHAI-TIME', 'SUNDOWNER'],
        'indian': ['MASALA', 'GARAM', 'TANDOORI', 'BIRYANI', 'CURRY', 'NAAN', 'ROTI', 'DOSA', 'IDLI', 'SAMOSA', 'PAKORA', 'GULAB-JAMUN', 'JALEBI', 'RASGULLA']
    }
    
    # 🎯 Achievement-based triggers
    achievement_toppings = []
    if calculation_data:
        revenue = calculation_data.get('revenue', 0)
        platform_fee = calculation_data.get('platform_fee', 0)
        margin = calculation_data.get('margin', '0%')
        
        # High value achievement
        if revenue > 10000:
            achievement_toppings.append('MONEY-BAGS')
        elif revenue > 5000:
            achievement_toppings.append('RICH-BOY')
        elif revenue > 1000:
            achievement_toppings.append('BUDGET-BOSS')
        
        # Platform fee achievement
        if platform_fee > 5000:
            achievement_toppings.append('PREMIUM-PLATFORM')
        elif platform_fee > 2000:
            achievement_toppings.append('STANDARD-PLATFORM')
        
        # Margin achievement
        try:
            margin_value = float(margin.replace('%', ''))
            if margin_value > 95:
                achievement_toppings.append('GOLDEN-RATIO')
            elif margin_value > 90:
                achievement_toppings.append('SILVER-STAR')
            elif margin_value > 80:
                achievement_toppings.append('BRONZE-BADGE')
        except:
            pass
        
        # 🍛 Indian-themed achievements
        if revenue > 50000:
            achievement_toppings.append('MAHARAJA')
        elif revenue > 25000:
            achievement_toppings.append('SARDAR')
        elif revenue > 15000:
            achievement_toppings.append('ZAMINDAR')
        
        # Spice level achievement (based on complexity)
        if calculation_data.get('complexity', 0) > 8:
            achievement_toppings.append('GARAM-MASALA')
        elif calculation_data.get('complexity', 0) > 5:
            achievement_toppings.append('MASALA')
    
    # 🕐 Time-based triggers
    current_hour = datetime.now().hour
    time_toppings = []
    if 6 <= current_hour < 11:
        time_toppings.append('BREAKFAST')
    elif 11 <= current_hour < 15:
        time_toppings.append('LUNCH')
    elif 15 <= current_hour < 19:
        time_toppings.append('DINNER')
    elif 19 <= current_hour < 23:
        time_toppings.append('EVENING')
    elif 23 <= current_hour or current_hour < 6:
        time_toppings.append('MIDNIGHT-SNACK')
    
    # 🍵 Indian time references
    if 7 <= current_hour < 9:
        time_toppings.append('CHAI-TIME')
    elif 16 <= current_hour < 18:
        time_toppings.append('EVENING-CHAI')
    elif 20 <= current_hour < 22:
        time_toppings.append('SUNDOWNER')
    
    # 🎲 Random pizza base (with higher chance for Indian options)
    base_toppings = random.choice(pizza_toppings['basic'])
    
    # 🏆 Combine all toppings
    all_toppings = [base_toppings] + achievement_toppings + time_toppings
    
    # 🍛 Add some Indian flair - 30% chance to include an Indian topping
    if random.random() < 0.3:
        indian_topping = random.choice(pizza_toppings['indian'])
        all_toppings.append(indian_topping)
    
    selected_toppings = random.sample(all_toppings, min(3, len(all_toppings)))
    
    # 🍕 Generate the fun ID
    pizza_id = '-'.join(selected_toppings)
    
    # Add some random numbers and letters for uniqueness
    random_part = ''.join(random.choices('0123456789ABCDEF', k=8))
    
    return f"{pizza_id}-{random_part}"

# 🎉 Easter egg trigger function
def should_trigger_easter_egg():
    """Randomly decide if we should use the easter egg ID (80% chance for testing, can be reduced later)"""
    import random
    return random.random() < 0.8  # 80% chance - increased for testing Indian toppings

def calculate_safe_overage_price(rate_card_price, meta_cost, markup_multiplier=1.2):
    """
    Calculate overage price ensuring it's never lower than the rate card price.
    
    Args:
        rate_card_price: Base rate card price
        meta_cost: Additional meta cost
        markup_multiplier: Multiplier for overage (default 1.2 = 20% markup)
    
    Returns:
        float: Safe overage price that's always >= rate card price
    """
    base_cost = rate_card_price + meta_cost
    overage_price = base_cost * markup_multiplier
    
    # Ensure overage price is never lower than rate card price
    safe_overage_price = max(overage_price, rate_card_price)
    
    return round(safe_overage_price, 4)

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Needed for session

# SOW beta allowlist – only these emails see the Generate SOW (beta) option
SOW_BETA_EMAILS = {
    'adwit.sharma@gupshup.io',
    'ankit.kanwara@gupshup.io',
    'gargi.upadhyay@gupshup.io',
    'kathyayani.nayak@gupshup.io',
    'mauricio.martins@gupshup.io',
    'mridul.kumawat@gupshup.io',
    'nikhil.sharma@knowlarity.com',
    'nikhil.sharma@gupshup.io',
    'purusottam.singh@gupshup.io',
    'siddharth.singh@gupshup.io',
    'yashas.reddy@gupshup.io',
}

# Test log to verify logging is working
print("=== FLASK APP STARTING UP ===", file=sys.stderr, flush=True)
print(f"Python version: {sys.version}", file=sys.stderr, flush=True)
print("=== END STARTUP LOG ===", file=sys.stderr, flush=True)

# Set up proper logging for Railway/gunicorn
import logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stderr)
    ]
)
logger = logging.getLogger(__name__)
logger.info("Flask app logger initialized")
# --- Session helpers ---
def clear_calc_session(preserve_auth: bool = True):
    """Clear all calculation-related session keys. Optionally preserve auth."""
    keys_to_clear = [
        'chosen_platform_fee', 'pricing_inputs', 'rate_card_platform_fee', 'results',
        'selected_components', 'user_selections', 'inclusions', 'calculation_id',
        'bundle_choice', 'manday_rates', 'inputs', 'voice_pricing'
    ]
    for k in keys_to_clear:
        if k in session:
            session.pop(k)
    if not preserve_auth and 'authenticated' in session:
        session.pop('authenticated')


# Health check endpoint - moved to top for immediate availability
@app.route('/health')
def health_check():
    logger.info("Health check endpoint accessed")
    print("HEALTH CHECK ACCESSED", file=sys.stderr, flush=True)
    return "OK", 200

# Root route for basic connectivity test
@app.route('/ping')
def ping():
    return "pong", 200

def initialize_inclusions():
    """
    Returns a dictionary of all possible inclusions for each feature/tier.
    Only the inclusions for the highest/most specific tier in each category should be shown.
    """
    inclusions = {
         'Platform Fee Used for Margin Calculation': [
            'Journey Builder Lite',
            'Campaign Manager',
            'CTWA - (Meta, Tiktok)',
            'Agent Assist < 20 Agents',
            'personalize lite upto 1 million records - no advanced events',
            '80 TPS',
            '1 manday/month maintenance'
        ],
        'Personalize Load Lite': [
            'personalize lite upto 1 million records - no advanced events'
        ],
        'Human Agents <20': [
            'Agent Assist < 20 Agents'
        ],
        'Human Agents 20+': [
            'Agent Assist 20-50 agents'
        ],
        'Human Agents 50+': [
            'Agent Assist 50-100 agents'
        ],
        'Human Agents 100+': [
            'Agent Assist 100+ agents, advanced routing rules'
        ],
        'BFSI Tier 1': [
            'Auditing to be stored for XX days for JB PRO+Flows',
            'Conversational Data Encryption, Logging,Auditing, Purging (Data and Logs)'
        ],
        'BFSI Tier 2': [
            'Auditing to be stored for XX days for JB PRO+Flows',
            'Conversational Data Encryption, Logging,Auditing, Purging (Data and Logs) on Agent Assist',
            'Conversational Data Encryption, Logging,Auditing,Purging (Data and Logs) on AI Agents',
            'Encryption, Auditing,Logging on Campaign Manager'
        ],
        'BFSI Tier 3': [
            'Auditing to be stored for XX days for JB PRO+Flows',
            'Conversational Data Encryption, Logging,Auditing, Purging (Data and Logs) on Agent Assist',
            'Conversational Data Encryption, Logging,Auditing,Purging (Data and Logs) on AI Agents',
            'Encryption, Auditing,Logging on Campaign Manager',
            'Data Encryption, Logging, Auditing, Purging (Logs) on Reatargetting and Personalize'
        ],
        'Personalize Load Standard': [
            'Standard - upto 5 million records, no external events supported'
        ],
        'Personalize Load Advanced': [
            'Advanced - upto 10 million records, external events supported'
        ],
        'AI Module Yes': [
            'AI- Workspace Configuration and Data retarining UI'
        ],
        'Increased TPS 250': [
            '250 - Upto 250 messages per second'
        ],
        'Increased TPS 1000': [
            '1000 - upto 1000 messages per second'
        ],
        'Smart CPaaS Yes': [
            'Smart CPaaS - Intelligent failover between channels'
        ]
    }
    return inclusions

def calculate_platform_fee(country, bfsi_tier, personalize_load, human_agents, ai_module, smart_cpaas, increased_tps='NA'):
    """
    Calculates the platform fee based on country and selected options.
    Returns (fee, currency).
    Now uses PLATFORM_PRICING_GUIDANCE for all values.
    """
    guidance = PLATFORM_PRICING_GUIDANCE.get(country, PLATFORM_PRICING_GUIDANCE['APAC'])
    fee = guidance['minimum']
    currency = 'INR' if country == 'India' else 'USD'
    # BFSI Tiers
    if bfsi_tier == 'Tier 1':
        fee += guidance.get('BFSI_Tier_1', 0)
    elif bfsi_tier == 'Tier 2':
        fee += guidance.get('BFSI_Tier_2', 0)
    elif bfsi_tier == 'Tier 3':
        fee += guidance.get('BFSI_Tier_3', 0)
    # TPS
    if increased_tps == '250':
        fee += guidance.get('TPS_250', 0)
    elif increased_tps == '1000':
        fee += guidance.get('TPS_1000', 0)
    # Personalize Load
    if personalize_load == 'Standard':
        fee += guidance.get('Personalize_Standard', 0)
    elif personalize_load in ['Advanced', 'Pro']:
        fee += guidance.get('Personalize_Pro', 0)
    # Agent Assist (Human Agents)
    if human_agents == '20+':
        fee += guidance.get('Agent_Assist_20_50', 0)
    elif human_agents == '50+':
        fee += guidance.get('Agent_Assist_50_100', 0)
    elif human_agents == '100+':
        fee += guidance.get('Agent_Assist_100_plus', 0)
    # AI Module
    if ai_module == 'Yes':
        fee += guidance.get('AI_Module', 0)
    # Smart CPaaS
    if smart_cpaas == 'Yes':
        fee += guidance.get('Smart_CPaaS', 0)
    return fee, currency

# Custom Jinja2 filter for number formatting
@app.template_filter('smart_format')
def smart_format_filter(value):
    """Format numbers: no decimals if zero, up to 4 decimals if non-zero"""
    try:
        if value is None:
            return ''
        val = float(value)
        # Check if the number is a whole number
        if val == int(val):
            return f"{int(val):,}"
        else:
            # Show up to 4 decimal places, but remove trailing zeros
            s = f"{val:.4f}"
            return f"{s.rstrip('0').rstrip('.'):,}"
    except (ValueError, TypeError):
        return str(value) if value is not None else ''

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'postgresql://postgres:prdeuXwtBzpLZaOGpxgRspfjfLNEQrys@gondola.proxy.rlwy.net:25504/railway')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
migrate = Migrate(app, db)


def record_funnel_event(step, inputs=None, profile=None):
    """
    Safely record a step-level funnel event for later analytics.
    Never breaks user flow if anything goes wrong.
    """
    try:
        inputs = inputs or session.get('inputs') or {}
        profile = profile or session.get('profile') or {}
        route = 'bundle' if float(inputs.get('committed_amount', 0) or 0) > 0 else 'volumes'
        ev = FunnelEvent(
            user_email=(profile or {}).get('email'),
            calculation_id=session.get('calculation_id'),
            step=step,
            route=route,
            country=inputs.get('country'),
            region=inputs.get('region'),
        )
        db.session.add(ev)
        db.session.commit()
    except Exception:
        logger.exception("Failed to record funnel event for step '%s'", step)

# Example Analytics model (expand as needed)
class Analytics(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    calculation_id = db.Column(db.String(64), unique=False, nullable=True)  # Transaction ID for each calculation
    timestamp = db.Column(db.DateTime, nullable=False)
    user_name = db.Column(db.String(128))
    user_email = db.Column(db.String(256), nullable=True)
    country = db.Column(db.String(64))
    region = db.Column(db.String(64))  # Added region column for region-specific analytics
    platform_fee = db.Column(db.Float)
    ai_price = db.Column(db.Float)
    advanced_price = db.Column(db.Float)
    basic_marketing_price = db.Column(db.Float)
    basic_utility_price = db.Column(db.Float)
    currency = db.Column(db.String(8))
    # New fields for advanced analytics
    ai_rate_card_price = db.Column(db.Float)
    advanced_rate_card_price = db.Column(db.Float)
    basic_marketing_rate_card_price = db.Column(db.Float)
    basic_utility_rate_card_price = db.Column(db.Float)
    ai_volume = db.Column(db.Float)
    advanced_volume = db.Column(db.Float)
    basic_marketing_volume = db.Column(db.Float)
    basic_utility_volume = db.Column(db.Float)
    # New fields for user-submitted manday rates
    bot_ui_manday_rate = db.Column(db.Float)
    custom_ai_manday_rate = db.Column(db.Float)
    # New fields for manday counts
    bot_ui_mandays = db.Column(db.Float)
    custom_ai_mandays = db.Column(db.Float)
    committed_amount = db.Column(db.Float, nullable=True)  # New column for message bundle amount
    # Voice notes fields
    voice_notes_price = db.Column(db.String(8))  # "Yes" or "No"
    voice_notes_model = db.Column(db.String(64))  # Model selected
    voice_notes_rate = db.Column(db.Float)  # Rate per minute
    # AI agent selection fields
    ai_agent_model = db.Column(db.String(128), nullable=True)
    ai_agent_complexity = db.Column(db.String(32), nullable=True)
    # Add more fields as needed
    calculation_route = db.Column(db.String(16))  # "volumes" or "bundle"
    # --- Voice channel analytics (nullable to remain backward compatible) ---
    channel_type = db.Column(db.String(32), nullable=True)  # text_only, voice_only, text_voice
    voice_mandays = db.Column(db.Float, nullable=True)
    voice_dev_cost = db.Column(db.Float, nullable=True)
    voice_platform_fee = db.Column(db.Float, nullable=True)
    whatsapp_setup_fee = db.Column(db.Float, nullable=True)
    pstn_inbound_ai_minutes = db.Column(db.Float, nullable=True)
    pstn_inbound_committed = db.Column(db.Float, nullable=True)
    pstn_outbound_ai_minutes = db.Column(db.Float, nullable=True)
    pstn_outbound_committed = db.Column(db.Float, nullable=True)
    pstn_manual_minutes = db.Column(db.Float, nullable=True)
    pstn_manual_committed = db.Column(db.Float, nullable=True)
    whatsapp_voice_outbound_minutes = db.Column(db.Float, nullable=True)
    whatsapp_voice_inbound_minutes = db.Column(db.Float, nullable=True)
    voice_cost_pstn_inbound = db.Column(db.Float, nullable=True)
    voice_cost_pstn_outbound = db.Column(db.Float, nullable=True)
    voice_cost_pstn_manual = db.Column(db.Float, nullable=True)
    voice_cost_wa_outbound = db.Column(db.Float, nullable=True)
    voice_cost_wa_inbound = db.Column(db.Float, nullable=True)
    voice_total_cost = db.Column(db.Float, nullable=True)
    # --- SOW funnel analytics ---
    sow_generate_clicked = db.Column(db.Boolean, nullable=True, default=False)
    sow_downloaded = db.Column(db.Boolean, nullable=True, default=False)


class FunnelEvent(db.Model):
    """
    Lightweight step-level funnel tracking for the calculator.
    Each row represents a user hitting a specific step in the flow.
    """
    __tablename__ = 'funnel_events'

    id = db.Column(db.Integer, primary_key=True)
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    user_email = db.Column(db.String(256), nullable=True)
    calculation_id = db.Column(db.String(64), nullable=True)
    step = db.Column(db.String(32), nullable=False)  # volumes, prices, results, sow_details, sow_download
    route = db.Column(db.String(16), nullable=True)  # volumes or bundle
    country = db.Column(db.String(64), nullable=True)
    region = db.Column(db.String(64), nullable=True)

# os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'  # For local testing only

# Country to currency symbol mapping
COUNTRY_CURRENCY = {
    'India': '₹',
    'MENA': '$',  # USD for MENA
    'LATAM': '$',
    'Africa': '$',
    'Europe': '$',  # Use USD for Europe
    'APAC': '$',
    'Rest of the World': '$',  # For historical data compatibility
}

SECRET_ANALYTICS_KEYWORD = "letmein123"

# Add this near the top
CALC_PASSWORD = os.environ.get('CALC_PASSWORD', 'gup$hup.i0')  # Set your password here or via env var

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        password = request.form.get('password', '')
        if password == CALC_PASSWORD:
            session['authenticated'] = True
            return redirect(url_for('profile_email'))
        else:
            flash('Incorrect password. Please try again.', 'error')
    else:
        # Reset session fully on GET to handle relogin or stale sessions
        try:
            session.clear()
        except Exception:
            pass
    return render_template('login.html')


@app.route('/profile-email', methods=['GET', 'POST'])
def profile_email():
    """
    Lightweight email capture step after password login.
    If we can find an existing profile for this email, skip profile page and go to volumes.
    Otherwise, show full profile page next.
    """
    if not session.get('authenticated'):
        return redirect(url_for('login'))

    from pricing_config import get_default_location_for_email

    profile = session.get('profile') or {}

    if request.method == 'POST':
        email = (request.form.get('user_email') or '').strip().lower()
        if not email:
            flash('Please enter a valid work email.', 'error')
            return render_template('profile_email.html', profile=profile)

        # Start building/refreshing profile with email
        profile['email'] = email

        # Try to infer defaults from mapping
        mapped = get_default_location_for_email(email)
        if mapped:
            country, region = mapped
            profile.setdefault('country', country)
            if region:
                profile.setdefault('region', region)

        # Try to hydrate from last Analytics entry for this email (pseudo-profile)
        existing = Analytics.query.filter_by(user_email=email).order_by(Analytics.timestamp.desc()).first()
        if existing:
            if existing.user_name:
                profile.setdefault('name', existing.user_name)
            if existing.country:
                profile.setdefault('country', existing.country)
            if existing.region:
                profile.setdefault('region', existing.region)

        session['profile'] = profile

        # Decide next step:
        # If we have at least name + country, treat profile as existing and go to volumes.
        if profile.get('name') and profile.get('country'):
            return redirect(url_for('index', step='volumes'))
        # Otherwise collect remaining details on profile page
        return redirect(url_for('index', step='profile'))

    # GET: just show email field, prefilled if we know it
    return render_template('profile_email.html', profile=profile)

@app.route('/start-over', methods=['GET'])
def start_over():
    """Explicit route to wipe calculation state and return to the start."""
    clear_calc_session(preserve_auth=True)
    flash('Started a fresh calculation.', 'success')
    return redirect(url_for('index'))

@app.route('/', methods=['GET', 'POST'])
def index():
    if not session.get('authenticated'):
        return redirect(url_for('login'))
    """
    Main route for the pricing calculator. Handles all user steps (volumes, prices, bundle, results).
    Manages session data, input validation, pricing logic, and inclusions logic.
    """
    print("DEBUG: session at start of request:", dict(session), file=sys.stderr, flush=True)
    # Determine current step; default to 'volumes' for backward compatibility.
    # We will explicitly route to 'profile' when needed below.
    step = request.form.get('step', request.args.get('step', 'volumes'))
    
    # If no form data (page refresh) and we have results in session, assume results step
    if not request.form and session.get('results') and session.get('inputs'):
        step = 'results'
    
    print("\n--- DEBUG ---", file=sys.stderr, flush=True)
    print("Form data:", dict(request.form), file=sys.stderr, flush=True)
    print("Session data:", dict(session), file=sys.stderr, flush=True)
    print("Detected step:", step, file=sys.stderr, flush=True)
    
    # Generate calculation_id only when user moves to page 2 (prices or bundle step)
    # For all other steps, preserve existing calculation_id from session
    if step in ['prices', 'bundle'] and not session.get('calculation_id'):
        # 🍕 Easter egg chance for calculation ID (only for new calculations)
        if should_trigger_easter_egg():
            calculation_id = generate_pizza_easter_egg_id()
            print(f"🍕 PIZZA EASTER EGG TRIGGERED! ID: {calculation_id}", file=sys.stderr, flush=True)
        else:
            calculation_id = str(uuid.uuid4())
        session['calculation_id'] = calculation_id
        print(f"DEBUG: New calculation started on {step} step. Calculation ID: {calculation_id}", file=sys.stderr, flush=True)
    else:
        # For all other cases, preserve the calculation_id from session
        calculation_id = session.get('calculation_id', str(uuid.uuid4()))
        print(f"DEBUG: Using existing calculation ID: {calculation_id}", file=sys.stderr, flush=True)
    
    print("Current step:", step, file=sys.stderr, flush=True)
    print("--- END DEBUG ---\n", file=sys.stderr, flush=True)
    results = None
    currency_symbol = None

    min_fees = {country: data['minimum'] for country, data in PLATFORM_PRICING_GUIDANCE.items()}

    # ---------------------------------------------------------------------
    # Profile step – still available but no longer forced before volumes.
    # We now primarily capture profile data via email + fields on Volumes.
    # ---------------------------------------------------------------------
    if step == 'profile':
        profile = session.get('profile', {}) or {}
        if request.method == 'POST':
            name = (request.form.get('user_name') or '').strip()
            email = (request.form.get('user_email') or '').strip()
            country = (request.form.get('country') or '').strip() or 'India'
            region = (request.form.get('region') or '').strip()

            # Basic email validation – allow blank, but if present require simple syntax
            if email and ('@' not in email or '.' not in email.split('@')[-1]):
                flash('Please enter a valid work email address.', 'error')
                profile = {
                    'name': name,
                    'email': email,
                    'country': country,
                    'region': region,
                }
                return render_template(
                    'index.html',
                    step='profile',
                    profile=profile,
                    calculation_id=calculation_id,
                    min_fees=min_fees,
                )

            # Auto-prefill country/region if mapping exists and user did not override
            mapped = get_default_location_for_email(email) if email else None
            if mapped:
                mapped_country, mapped_region = mapped
                if not request.form.get('country'):
                    country = mapped_country
                if not request.form.get('region') and mapped_region:
                    region = mapped_region

            profile = {
                'name': name,
                'email': email,
                'country': country,
                'region': region,
            }
            session['profile'] = profile

            # Mirror key fields into session['inputs'] so existing code keeps working
            inputs = session.get('inputs', {}) or {}
            if name:
                inputs['user_name'] = name
            if country:
                inputs['country'] = country
            if region is not None:
                inputs['region'] = region
            session['inputs'] = inputs

            # Continue to volumes step
            return redirect(url_for('index', step='volumes'))

        # GET: show profile form, prefilling from existing profile/inputs
        inputs = session.get('inputs', {}) or {}
        profile.setdefault('name', inputs.get('user_name', ''))
        profile.setdefault('email', '')
        profile.setdefault('country', inputs.get('country', 'India'))
        profile.setdefault('region', inputs.get('region', ''))
        return render_template(
            'index.html',
            step='profile',
            profile=profile,
            calculation_id=calculation_id,
            min_fees=min_fees,
        )

    # Defensive: ensure session data exists for edit actions
    if step == 'volumes' and request.method == 'POST':
        # Clear previous session state for a new calculation
        for key in ['chosen_platform_fee', 'pricing_inputs', 'rate_card_platform_fee', 'results', 'selected_components', 'user_selections', 'inclusions', 'calculation_id']:
            if key in session:
                session.pop(key)
        # Calculation ID will be generated when user moves to page 2 (prices or bundle)
        user_name = request.form.get('user_name', '')
        user_email = (request.form.get('user_email') or '').strip()
        # Step 1: User submitted volumes and platform fee options
        country = request.form['country'].strip()  # Always strip country
        region = request.form.get('region', '')
        # Ignore region for Europe and APAC
        if country in ['Europe', 'APAC']:
            region = ''
        # Set currency: Only use INR for India, all others use USD
        if country == 'India':
            currency = 'INR'
        else:
            currency = 'USD'
        # Use dev_location only for India
        dev_location = request.form.get('dev_location', 'India').strip() if country == 'India' else None
        print(f"[app.py] After country/currency logic: country='{country}', currency='{currency}', dev_location='{dev_location}'", file=sys.stderr, flush=True)
        def parse_volume(val):
            try:
                return float(val.replace(',', '')) if val and str(val).strip() else 0.0
            except Exception:
                return 0.0
        ai_volume = parse_volume(request.form.get('ai_volume', ''))
        advanced_volume = parse_volume(request.form.get('advanced_volume', ''))
        basic_marketing_volume = parse_volume(request.form.get('basic_marketing_volume', ''))
        basic_utility_volume = parse_volume(request.form.get('basic_utility_volume', ''))
        # Debug: Log parsed volumes
        print(f"DEBUG: Parsed volumes - ai: {ai_volume}, advanced: {advanced_volume}, marketing: {basic_marketing_volume}, utility: {basic_utility_volume}", file=sys.stderr, flush=True)
        bfsi_tier = request.form.get('bfsi_tier', 'NA')
        personalize_load = request.form.get('personalize_load', 'NA')
        human_agents = request.form.get('human_agents', 'NA')
        ai_module = request.form.get('ai_module', 'NA')
        smart_cpaas = request.form.get('smart_cpaas', 'No')
        increased_tps = request.form.get('increased_tps', 'NA')
        # --- One Time Dev Activity Fields ---
        onboarding_price = request.form.get('onboarding_price', 'No')
        ux_price = request.form.get('ux_price', 'No')
        testing_qa_price = request.form.get('testing_qa_price', 'No')
        aa_setup_price = request.form.get('aa_setup_price', 'No')
        num_apis_price = request.form.get('num_apis_price', '0')
        num_journeys_price = request.form.get('num_journeys_price', '0')
        num_ai_workspace_commerce_price = request.form.get('num_ai_workspace_commerce_price', '0')
        num_ai_workspace_faq_price = request.form.get('num_ai_workspace_faq_price', '0')
        # If any AI workspace is selected, AI Module must be Yes
        try:
            ws_commerce_count = int(str(num_ai_workspace_commerce_price).strip() or '0')
        except Exception:
            ws_commerce_count = 0
        try:
            ws_faq_count = int(str(num_ai_workspace_faq_price).strip() or '0')
        except Exception:
            ws_faq_count = 0
        if ws_commerce_count > 0 or ws_faq_count > 0:
            ai_module = 'Yes'
        # Voice notes fields
        voice_notes_price = request.form.get('voice_notes_price', 'No')
        voice_notes_model = request.form.get('voice_notes_model', '')
        # AI agent model & complexity (optional, may be required for AI module flows)
        ai_agent_model = request.form.get('ai_agent_model', '')
        ai_agent_complexity = request.form.get('ai_agent_complexity', 'regular')
        # Voice channel fields
        channel_type = request.form.get('channel_type', 'text_only')
        # Enforce: Voice supported only for India
        if country != 'India' and channel_type in ['voice_only', 'text_voice']:
            channel_type = 'text_only'
        def pv(name):
            return parse_volume(request.form.get(name, ''))
        pstn_inbound_ai_minutes = pv('pstn_inbound_ai_minutes')
        pstn_inbound_committed = pv('pstn_inbound_committed')
        pstn_outbound_ai_minutes = pv('pstn_outbound_ai_minutes')
        pstn_outbound_committed = pv('pstn_outbound_committed')
        pstn_manual_minutes = pv('pstn_manual_minutes')
        pstn_manual_committed = pv('pstn_manual_committed')
        whatsapp_voice_outbound_minutes = pv('whatsapp_voice_outbound_minutes')
        whatsapp_voice_inbound_minutes = pv('whatsapp_voice_inbound_minutes')
        voice_ai_enabled = request.form.get('voice_ai_enabled', 'No')
        num_voice_journeys = request.form.get('num_voice_journeys', '0')
        num_voice_apis = request.form.get('num_voice_apis', '0')
        num_additional_voice_languages = request.form.get('num_additional_voice_languages', '0')
        agent_handover_pstn = request.form.get('agent_handover_pstn', 'None')
        whatsapp_voice_platform = request.form.get('whatsapp_voice_platform', 'None')
        virtual_number_required = request.form.get('virtual_number_required', 'No')
        platform_fee, fee_currency = calculate_platform_fee(country, bfsi_tier, personalize_load, human_agents, ai_module, smart_cpaas, increased_tps)
        currency_symbol = COUNTRY_CURRENCY.get(country, '$')
        # Always update session['inputs'] with latest form data
        session['inputs'] = {
            'user_name': user_name,
            'user_email': user_email,
            'country': country,
            'region': region,
            'ai_volume': ai_volume,
            'advanced_volume': advanced_volume,
            'basic_marketing_volume': basic_marketing_volume,
            'basic_utility_volume': basic_utility_volume,
            'platform_fee': platform_fee,
            'bfsi_tier': bfsi_tier,
            'personalize_load': personalize_load,
            'human_agents': human_agents,
            'ai_module': ai_module,
            'ai_agent_model': ai_agent_model,
            'ai_agent_complexity': ai_agent_complexity,
            'smart_cpaas': smart_cpaas,
            'increased_tps': increased_tps,
            # One time dev fields
            'onboarding_price': onboarding_price,
            'ux_price': ux_price,
            'testing_qa_price': testing_qa_price,
            'aa_setup_price': aa_setup_price,
            'num_apis_price': num_apis_price,
            'num_journeys_price': num_journeys_price,
            'num_ai_workspace_commerce_price': num_ai_workspace_commerce_price,
            'num_ai_workspace_faq_price': num_ai_workspace_faq_price,
            # Voice notes fields
            'voice_notes_price': voice_notes_price,
            'voice_notes_model': voice_notes_model,
            # Voice channel fields
            'channel_type': channel_type,
            'voice_ai_enabled': voice_ai_enabled,
            'num_voice_journeys': num_voice_journeys,
            'num_voice_apis': num_voice_apis,
            'num_additional_voice_languages': num_additional_voice_languages,
            'agent_handover_pstn': agent_handover_pstn,
            'whatsapp_voice_platform': whatsapp_voice_platform,
            'virtual_number_required': virtual_number_required,
            'pstn_inbound_ai_minutes': pstn_inbound_ai_minutes,
            'pstn_inbound_committed': pstn_inbound_committed,
            'pstn_outbound_ai_minutes': pstn_outbound_ai_minutes,
            'pstn_outbound_committed': pstn_outbound_committed,
            'pstn_manual_minutes': pstn_manual_minutes,
            'pstn_manual_committed': pstn_manual_committed,
            'whatsapp_voice_outbound_minutes': whatsapp_voice_outbound_minutes,
            'whatsapp_voice_inbound_minutes': whatsapp_voice_inbound_minutes,
        }
        # Ensure profile is updated and mirrored into inputs
        profile = session.get('profile') or {}
        if user_email:
            profile['email'] = user_email
        if user_name:
            profile.setdefault('name', user_name)
        if country:
            profile.setdefault('country', country)
        if region is not None:
            profile.setdefault('region', region)
        session['profile'] = profile

        if profile.get('name'):
            session['inputs']['user_name'] = profile['name']
        if profile.get('email'):
            session['inputs']['user_email'] = profile['email']
        if profile.get('country'):
            session['inputs']['country'] = profile['country']
        if profile.get('region') is not None:
            session['inputs']['region'] = profile['region']
        print("DEBUG: session['inputs'] just set to:", session['inputs'], file=sys.stderr, flush=True)
        # Only route to bundle if all text volumes are zero AND this is not a voice-only scenario
        total_voice_minutes = (
            pstn_inbound_ai_minutes + pstn_outbound_ai_minutes + pstn_manual_minutes +
            whatsapp_voice_outbound_minutes + whatsapp_voice_inbound_minutes
        )
        if all(float(v) == 0.0 for v in [ai_volume, advanced_volume, basic_marketing_volume, basic_utility_volume]) and channel_type != 'voice_only' and total_voice_minutes == 0:
            currency_symbol = COUNTRY_CURRENCY.get(country, '$')
            return render_template('index.html', step='bundle', currency_symbol=currency_symbol, inputs=session.get('inputs', {}), platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees)
        # Suggest prices
        def is_zero(val):
            try:
                return float(val) == 0.0
            except Exception:
                return True
        # --- Voice rate card (read-only) for Prices page ---
        voice_rate_card = None
        try:
            chan = session.get('inputs', {}).get('channel_type', 'text_only')
            if country == 'India' and chan in ['voice_only', 'text_voice']:
                from pricing_config import PSTN_CALLING_CHARGES, get_whatsapp_voice_rate
                wa_out = float(session['inputs'].get('whatsapp_voice_outbound_minutes', 0) or 0)
                wa_in = float(session['inputs'].get('whatsapp_voice_inbound_minutes', 0) or 0)
                total_wa_min = wa_out + wa_in
                voice_rate_card = {
                    'pstn': {
                        'inbound': PSTN_CALLING_CHARGES['inbound_ai'],
                        'outbound': PSTN_CALLING_CHARGES['outbound_ai'],
                        'manual': PSTN_CALLING_CHARGES['manual_c2c'],
                    },
                    'whatsapp': {
                        'outbound': get_whatsapp_voice_rate(country, total_wa_min, 'outbound'),
                        'inbound': get_whatsapp_voice_rate(country, total_wa_min, 'inbound'),
                        'total_minutes': total_wa_min,
                    }
                }
        except Exception:
            voice_rate_card = None

        # --- AI pricing: combine tier-based markup with optional model + complexity ---
        base_ai_markup = get_suggested_price(country, 'ai', ai_volume) if not is_zero(ai_volume) else get_lowest_tier_price(country, 'ai')
        ai_model = session['inputs'].get('ai_agent_model', '')
        ai_complexity = session['inputs'].get('ai_agent_complexity', 'regular')
        ai_components = compute_ai_price_components(country, ai_model, ai_complexity, base_ai_markup)

        suggested_prices = {
            'ai_price': ai_components['markup'],
            'advanced_price': get_suggested_price(country, 'advanced', advanced_volume) if not is_zero(advanced_volume) else get_lowest_tier_price(country, 'advanced'),
            'basic_marketing_price': get_suggested_price(country, 'basic_marketing', basic_marketing_volume) if not is_zero(basic_marketing_volume) else get_lowest_tier_price(country, 'basic_marketing'),
            'basic_utility_price': get_suggested_price(country, 'basic_utility', basic_utility_volume) if not is_zero(basic_utility_volume) else get_committed_amount_rate_for_volume(country, 'basic_utility', 0),
            'voice_notes_rate': get_voice_notes_price(country, session.get('inputs', {}).get('voice_notes_model', '')) if session.get('inputs', {}).get('voice_notes_price') == 'Yes' else 0.0,
        }
        suggested_prices = patch_suggested_prices(suggested_prices, session.get('inputs', {}))
        return render_template('index.html', step='prices', suggested=suggested_prices, inputs=session.get('inputs', {}), currency_symbol=currency_symbol, platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees, voice_rate_card=voice_rate_card)

    elif step == 'prices' and request.method == 'POST':
        print('HANDLER: Entered prices POST step', file=sys.stderr, flush=True)
        inputs = session.get('inputs', {})
        if not inputs:
            print('HANDLER: No inputs in session, redirecting to index', file=sys.stderr, flush=True)
            flash('Session expired or missing. Please start again.', 'error')
            return redirect(url_for('index'))
        def parse_price(val):
            try:
                return float(val.replace(',', '')) if val and str(val).strip() else None
            except Exception:
                return None
        ai_price = parse_price(request.form.get('ai_price', ''))
        advanced_price = parse_price(request.form.get('advanced_price', ''))
        basic_marketing_price = parse_price(request.form.get('basic_marketing_price', ''))
        basic_utility_price = parse_price(request.form.get('basic_utility_price', ''))
        platform_fee = parse_price(request.form.get('platform_fee', '')) or 0.0
        # Parse manday rates from form and save to session
        bot_ui_manday_rate = parse_price(request.form.get('bot_ui_manday_rate', ''))
        custom_ai_manday_rate = parse_price(request.form.get('custom_ai_manday_rate', ''))
        session['manday_rates'] = {
            'bot_ui': bot_ui_manday_rate,
            'custom_ai': custom_ai_manday_rate
        }
        # Capture optional Voice rate overrides
        voice_rate_overrides = {}
        for k in ['vr_pstn_in_bundled','vr_pstn_in_overage','vr_pstn_out_bundled','vr_pstn_out_overage','vr_pstn_manual_bundled','vr_pstn_manual_overage','vr_wa_out_per_min','vr_wa_in_per_min']:
            v = request.form.get(k, '')
            try:
                voice_rate_overrides[k] = float(v.replace(',', '')) if v and str(v).strip() else None
            except Exception:
                voice_rate_overrides[k] = None
        # Merge voice overrides into inputs for downstream calculation
        for k, v in voice_rate_overrides.items():
            if v is not None:
                inputs[k] = v

        # Get suggested prices for validation
        country = inputs.get('country', 'India')
        ai_volume = float(inputs.get('ai_volume', 0) or 0)
        advanced_volume = float(inputs.get('advanced_volume', 0) or 0)
        basic_marketing_volume = float(inputs.get('basic_marketing_volume', 0) or 0)
        basic_utility_volume = float(inputs.get('basic_utility_volume', 0) or 0)
        base_ai_markup = get_suggested_price(country, 'ai', ai_volume)
        ai_model = inputs.get('ai_agent_model', '')
        ai_complexity = inputs.get('ai_agent_complexity', 'regular')
        ai_components = compute_ai_price_components(country, ai_model, ai_complexity, base_ai_markup)
        suggested_ai = ai_components['markup']
        suggested_advanced = get_suggested_price(country, 'advanced', advanced_volume)
        suggested_marketing = get_suggested_price(country, 'basic_marketing', basic_marketing_volume)
        suggested_utility = get_suggested_price(country, 'basic_utility', basic_utility_volume)
        # Calculate rate card platform fee for discount check
        rate_card_platform_fee, _ = calculate_platform_fee(
            country,
            inputs.get('bfsi_tier', 'NA'),
            inputs.get('personalize_load', 'NA'),
            inputs.get('human_agents', 'NA'),
            inputs.get('ai_module', 'NA'),
            inputs.get('smart_cpaas', 'No'),
            inputs.get('increased_tps', 'NA')
        )
        discount_errors = []
        # Discount checks: allow up to 90% discount for basic marketing, 70% for all others
        if ai_price is not None and (suggested_ai or 0) and ai_price < 0.3 * (suggested_ai or 0):
            discount_errors.append("AI Message price is less than 30% of the rate card.")
        if advanced_price is not None and (suggested_advanced or 0) and advanced_price < 0.3 * (suggested_advanced or 0):
            discount_errors.append("Advanced Message price is less than 30% of the rate card.")
        if basic_marketing_price is not None and (suggested_marketing or 0) and basic_marketing_price < 0.1 * (suggested_marketing or 0):
            discount_errors.append("Basic Marketing Message price is less than 10% of the rate card.")
        if basic_utility_price is not None and (suggested_utility or 0) and basic_utility_price < 0.3 * (suggested_utility or 0):
            discount_errors.append("Basic Utility/Authentication Message price is less than 30% of the rate card.")
        # Platform fee discount check
        if platform_fee < 0.3 * rate_card_platform_fee:
            discount_errors.append("Platform Fee is less than 30% of the rate card platform fee.")
        if discount_errors:
            print('HANDLER: Discount errors found, rendering prices page with errors', file=sys.stderr, flush=True)
            for msg in discount_errors:
                flash(msg, 'error')
            flash("Probability of deal desk rejection is high.", 'error')
            # Recompute voice rate card on error re-render
            voice_rate_card = None
            try:
                chan = inputs.get('channel_type', 'text_only')
                if country == 'India' and chan in ['voice_only', 'text_voice']:
                    from pricing_config import PSTN_CALLING_CHARGES, get_whatsapp_voice_rate
                    wa_out = float(inputs.get('whatsapp_voice_outbound_minutes', 0) or 0)
                    wa_in = float(inputs.get('whatsapp_voice_inbound_minutes', 0) or 0)
                    total_wa_min = wa_out + wa_in
                    voice_rate_card = {
                        'pstn': {
                            'inbound': PSTN_CALLING_CHARGES['inbound_ai'],
                            'outbound': PSTN_CALLING_CHARGES['outbound_ai'],
                            'manual': PSTN_CALLING_CHARGES['manual_c2c'],
                        },
                        'whatsapp': {
                            'outbound': get_whatsapp_voice_rate(country, total_wa_min, 'outbound'),
                            'inbound': get_whatsapp_voice_rate(country, total_wa_min, 'inbound'),
                            'total_minutes': total_wa_min,
                        }
                    }
            except Exception:
                voice_rate_card = None

            suggested_prices = {
                'ai_price': suggested_ai,
                'advanced_price': suggested_advanced,
                'basic_marketing_price': suggested_marketing,
                'basic_utility_price': suggested_utility,
                'voice_notes_rate': get_voice_notes_price(country, inputs.get('voice_notes_model', '')) if inputs.get('voice_notes_price') == 'Yes' else 0.0,
            }
            currency_symbol = COUNTRY_CURRENCY.get(country, '$')
            # Re-render the pricing page with user input and error
            return render_template('index.html', step='prices', suggested=suggested_prices, inputs=inputs, currency_symbol=currency_symbol, platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees, voice_rate_card=voice_rate_card)
        print('HANDLER: No discount errors, continuing to results calculation', file=sys.stderr, flush=True)

        # Always recalculate platform fee before saving to session['pricing_inputs']
        calculated_platform_fee, fee_currency = calculate_platform_fee(
            country,
            inputs.get('bfsi_tier', 'NA'),
            inputs.get('personalize_load', 'NA'),
            inputs.get('human_agents', 'NA'),
            inputs.get('ai_module', 'NA'),
            inputs.get('smart_cpaas', 'No'),
            inputs.get('increased_tps', 'NA')
        )
        # Use user input if provided and valid, else use calculated
        user_platform_fee = request.form.get('platform_fee')
        if user_platform_fee and float(user_platform_fee) > 0:
            platform_fee = float(user_platform_fee)
        else:
            platform_fee = calculated_platform_fee
        # Overwrite any previous platform_fee value for this calculation
        session['pricing_inputs'] = {
            'ai_price': ai_price,
            'advanced_price': advanced_price,
            'basic_marketing_price': basic_marketing_price,
            'basic_utility_price': basic_utility_price,
            'platform_fee': platform_fee
        }

        currency_symbol = COUNTRY_CURRENCY.get(country, '$')
        # --- Restored logic for results calculation and rendering ---
        bundle_choice = 'No'
        session['bundle_choice'] = bundle_choice
        ai_price = session['pricing_inputs'].get('ai_price', 0)
        advanced_price = session['pricing_inputs'].get('advanced_price', 0)
        basic_marketing_price = session['pricing_inputs'].get('basic_marketing_price', 0)
        basic_utility_price = session['pricing_inputs'].get('basic_utility_price', 0)
        platform_fee = session['pricing_inputs'].get('platform_fee', 0)
        bundle_lines = []
        bundle_cost = 0
        for label, key, price in [
            ("AI Message", 'ai_volume', ai_price),
            ("Advanced Message", 'advanced_volume', advanced_price),
            ("Basic Marketing Message", 'basic_marketing_volume', basic_marketing_price),
            ("Basic Utility/Authentication Message", 'basic_utility_volume', basic_utility_price),
        ]:
            volume = float(inputs.get(key, 0))
            bundle_lines.append({
                'label': label,
                'volume': volume,
                'price': float(price)
            })
            bundle_cost += volume * float(price)
        total_bundle_price = float(platform_fee) + bundle_cost
        bundle_details = {
            'lines': bundle_lines,
            'bundle_cost': bundle_cost,
            'total_bundle_price': total_bundle_price,
            'inclusion_text': 'See table below for included volumes and overage prices.'
        }
        # Use one-time dev activity fields from form if present, else from session
        dev_fields = [
            'onboarding_price', 'ux_price', 'testing_qa_price', 'aa_setup_price',
            'num_apis_price', 'num_journeys_price', 'num_ai_workspace_commerce_price', 'num_ai_workspace_faq_price',
            'voice_notes_price', 'voice_notes_model', 'voice_notes_rate'
        ]
        patched_form = (request.form.copy() if hasattr(request.form, 'copy') else dict(request.form)) or {}
        for field in dev_fields:
            if field not in patched_form and field in inputs:
                patched_form[field] = inputs[field]
        for k in list(patched_form.keys()):
            if k.endswith('_price') or k.endswith('_rate') or k.endswith('_volume') or k.startswith('num_'):
                try:
                    patched_form[k] = str(patched_form[k]).replace(',', '')
                except Exception:
                    pass
        # --- Ensure manday_rates is always set and complete ---
        country = inputs.get('country', 'India')
        dev_location = inputs.get('dev_location', 'India')
        rates = COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])
        # Always reset manday_rates to backend defaults when country or dev_location changes
        if country == 'LATAM':
            default_bot_ui = float(rates['bot_ui'].get(dev_location, 0.0) or 0.0)
            default_custom_ai = float(rates['custom_ai'].get(dev_location, 0.0) or 0.0)
        else:
            default_bot_ui = float(rates.get('bot_ui', 0.0) or 0.0)
            default_custom_ai = float(rates.get('custom_ai', 0.0) or 0.0)
        manday_rates = {
            'bot_ui': default_bot_ui,
            'custom_ai': default_custom_ai,
            'default_bot_ui': default_bot_ui,
            'default_custom_ai': default_custom_ai
        }
        # Fill missing keys with defaults
        manday_rates['bot_ui'] = float(manday_rates.get('bot_ui', default_bot_ui))
        manday_rates['custom_ai'] = float(manday_rates.get('custom_ai', default_custom_ai))
        manday_rates['default_bot_ui'] = float(default_bot_ui)
        manday_rates['default_custom_ai'] = float(default_custom_ai)
        # Calculate discount percentages
        manday_rates['bot_ui_discount'] = round(100 * (default_bot_ui - manday_rates['bot_ui']) / default_bot_ui, 3) if default_bot_ui else 0
        manday_rates['custom_ai_discount'] = round(100 * (default_custom_ai - manday_rates['custom_ai']) / default_custom_ai, 3) if default_custom_ai else 0
        total_mandays = calculate_total_mandays(patched_form)
        # --- PATCH: Always .strip() the country and set dev_location only for India before dev cost calculation ---
        if 'country' in patched_form:
            patched_form['country'] = patched_form['country'].strip()
        else:
            patched_form['country'] = inputs.get('country', 'India').strip()
        if patched_form['country'] != 'India':
            patched_form['dev_location'] = None
        else:
            patched_form['dev_location'] = patched_form.get('dev_location', 'India').strip()
        total_dev_cost, dev_cost_currency, dev_cost_breakdown = calculate_total_manday_cost(patched_form, manday_rates)
        print(f"DEBUG: dev_cost_currency = {dev_cost_currency}, country = {country}", file=sys.stderr, flush=True)
        manday_breakdown = dev_cost_breakdown['mandays_breakdown']
        
        # Debug: Log inputs being used for calculation
        print(f"DEBUG: Calculation inputs - ai_volume: {inputs.get('ai_volume', 0)}, advanced_volume: {inputs.get('advanced_volume', 0)}, marketing_volume: {inputs.get('basic_marketing_volume', 0)}, utility_volume: {inputs.get('basic_utility_volume', 0)}", file=sys.stderr, flush=True)
        
        # Calculate pricing results (text channels)
        try:
            results = calculate_pricing(
                country=inputs.get('country', 'India'),
                ai_volume=float(inputs.get('ai_volume', 0) or 0),
                advanced_volume=float(inputs.get('advanced_volume', 0) or 0),
                basic_marketing_volume=float(inputs.get('basic_marketing_volume', 0) or 0),
                basic_utility_volume=float(inputs.get('basic_utility_volume', 0) or 0),
                platform_fee=float(platform_fee),
                ai_price=ai_price,
                advanced_price=advanced_price,
                basic_marketing_price=basic_marketing_price,
                basic_utility_price=basic_utility_price,
                voice_notes_rate=float(inputs.get('voice_notes_rate', 0) or 0) if inputs.get('voice_notes_price') == 'Yes' else None
            )
            print(f"DEBUG: Calculation successful, results: {results}", file=sys.stderr, flush=True)
        except Exception as e:
            print(f"DEBUG: Calculation failed with error: {e}", file=sys.stderr, flush=True)
            flash('Pricing calculation failed. Please check your inputs and try again.', 'error')
            suggested_prices = {
                'ai_price': suggested_ai,
                'advanced_price': suggested_advanced,
                'basic_marketing_price': suggested_marketing,
            'basic_utility_price': suggested_utility,
            'voice_notes_rate': get_voice_notes_price(country, inputs.get('voice_notes_model', '')) if inputs.get('voice_notes_price') == 'Yes' else 0.0,
        }
            currency_symbol = COUNTRY_CURRENCY.get(country, '$')
            return render_template('index.html', step='prices', suggested=suggested_prices, inputs=inputs, currency_symbol=currency_symbol, platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees)
        
        # Remove duplicate Committed Amount if present
        seen = set()
        unique_line_items = []
        if results and 'line_items' in results:
            for item in results['line_items']:
                key = (item.get('line_item'), item.get('chosen_price'), item.get('suggested_price'))
                if key not in seen:
                    unique_line_items.append(item)
                    seen.add(key)
            results['line_items'] = unique_line_items
        # Validate results
        if not results or 'line_items' not in results:
            print(f"DEBUG: Invalid results: {results}", file=sys.stderr, flush=True)
            flash('Pricing calculation failed. Please check your inputs and try again.', 'error')
            suggested_prices = {
                'ai_price': suggested_ai,
                'advanced_price': suggested_advanced,
                'basic_marketing_price': suggested_marketing,
            'basic_utility_price': suggested_utility,
            'voice_notes_rate': get_voice_notes_price(country, inputs.get('voice_notes_model', '')) if inputs.get('voice_notes_price') == 'Yes' else 0.0,
        }
            currency_symbol = COUNTRY_CURRENCY.get(country, '$')
            return render_template('index.html', step='prices', suggested=suggested_prices, inputs=inputs, currency_symbol=currency_symbol, platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees)
        print("PASSED results validation, about to render results page", file=sys.stderr, flush=True)
        try:
            # --- Ensure manday_rates is always set and complete ---
            country = inputs.get('country', 'India')
            dev_location = inputs.get('dev_location', 'India')
            rates = COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])
            # Always reset manday_rates to backend defaults when country or dev_location changes
            if country == 'LATAM':
                default_bot_ui = float(rates['bot_ui'].get(dev_location, 0.0) or 0.0)
                default_custom_ai = float(rates['custom_ai'].get(dev_location, 0.0) or 0.0)
            else:
                default_bot_ui = float(rates.get('bot_ui', 0.0) or 0.0)
                default_custom_ai = float(rates.get('custom_ai', 0.0) or 0.0)
            manday_rates = {
                'bot_ui': default_bot_ui,
                'custom_ai': default_custom_ai,
                'default_bot_ui': default_bot_ui,
                'default_custom_ai': default_custom_ai
            }
            # Fill missing keys with defaults
            manday_rates['bot_ui'] = float(manday_rates.get('bot_ui', default_bot_ui))
            manday_rates['custom_ai'] = float(manday_rates.get('custom_ai', default_custom_ai))
            manday_rates['default_bot_ui'] = float(default_bot_ui)
            manday_rates['default_custom_ai'] = float(default_custom_ai)
            # Calculate discount percentages
            manday_rates['bot_ui_discount'] = round(100 * (default_bot_ui - manday_rates['bot_ui']) / default_bot_ui, 3) if default_bot_ui else 0
            manday_rates['custom_ai_discount'] = round(100 * (default_custom_ai - manday_rates['custom_ai']) / default_custom_ai, 3) if default_custom_ai else 0
            total_mandays = calculate_total_mandays(patched_form)
            total_dev_cost, dev_cost_currency, dev_cost_breakdown = calculate_total_manday_cost(patched_form, manday_rates)
            print(f"DEBUG: dev_cost_currency = {dev_cost_currency}, country = {country}", file=sys.stderr, flush=True)
            manday_breakdown = dev_cost_breakdown['mandays_breakdown']
            # Debug: Log inputs being used for calculation
            print(f"DEBUG: Calculation inputs - ai_volume: {inputs.get('ai_volume', 0)}, advanced_volume: {inputs.get('advanced_volume', 0)}, marketing_volume: {inputs.get('basic_marketing_volume', 0)}, utility_volume: {inputs.get('basic_utility_volume', 0)}", file=sys.stderr, flush=True)
            # Calculate pricing results
            try:
                results = calculate_pricing(
                    country=inputs.get('country', 'India'),
                    ai_volume=float(inputs.get('ai_volume', 0) or 0),
                    advanced_volume=float(inputs.get('advanced_volume', 0) or 0),
                    basic_marketing_volume=float(inputs.get('basic_marketing_volume', 0) or 0),
                    basic_utility_volume=float(inputs.get('basic_utility_volume', 0) or 0),
                    platform_fee=float(platform_fee),
                    ai_price=ai_price,
                    advanced_price=advanced_price,
                    basic_marketing_price=basic_marketing_price,
                    basic_utility_price=basic_utility_price,
                    voice_notes_rate=float(inputs.get('voice_notes_rate', 0) or 0) if inputs.get('voice_notes_price') == 'Yes' else None
                )
                print(f"DEBUG: Calculation successful, results: {results}", file=sys.stderr, flush=True)
                # Voice pricing (if selected)
                channel_type = inputs.get('channel_type', 'text_only')
                has_text_ai = inputs.get('ai_module', 'NA') == 'Yes'
                if channel_type in ['voice_only', 'text_voice']:
                    from calculator import calculate_voice_pricing
                    voice_pricing = calculate_voice_pricing(inputs, country=inputs.get('country', 'India'), has_text_ai=has_text_ai)
                    session['voice_pricing'] = voice_pricing
                    results['voice_pricing'] = voice_pricing
            except Exception as e:
                print(f"DEBUG: Calculation failed with error: {e}", file=sys.stderr, flush=True)
                flash('Pricing calculation failed. Please check your inputs and try again.', 'error')
                suggested_prices = {
                    'ai_price': suggested_ai,
                    'advanced_price': suggested_advanced,
                    'basic_marketing_price': suggested_marketing,
            'basic_utility_price': suggested_utility,
            'voice_notes_rate': get_voice_notes_price(country, inputs.get('voice_notes_model', '')) if inputs.get('voice_notes_price') == 'Yes' else 0.0,
        }
                currency_symbol = COUNTRY_CURRENCY.get(country, '$')
                return render_template('index.html', step='prices', suggested=suggested_prices, inputs=inputs, currency_symbol=currency_symbol, platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees)
        except Exception as e:
            print(f"DEBUG: Error in manday/dev cost or pricing calculation: {e}", file=sys.stderr, flush=True)
            flash('Internal error during calculation. Please try again.', 'error')
            return render_template('index.html', step='prices', suggested={}, inputs=inputs, currency_symbol=currency_symbol, platform_fee=platform_fee, calculation_id=calculation_id, min_fees=min_fees)
        results['margin'] = results.get('margin', '')
        expected_invoice_amount = results.get('revenue', 0)
        chosen_platform_fee = float(platform_fee)
        rate_card_platform_fee, _ = calculate_platform_fee(
            inputs['country'],
            inputs.get('bfsi_tier', 'NA'),
            inputs.get('personalize_load', 'NA'),
            inputs.get('human_agents', 'NA'),
            inputs.get('ai_module', 'NA'),
            inputs.get('smart_cpaas', 'No'),
            inputs.get('increased_tps', 'NA')
        )
        user_selections = []
        if inputs.get('bfsi_tier', 'NA') not in ['NA', 'No']:
            user_selections.append(('BFSI Tier', inputs['bfsi_tier']))
        if inputs.get('personalize_load', 'NA') not in ['NA', 'No']:
            user_selections.append(('Personalize Load', inputs['personalize_load']))
        if inputs.get('human_agents', 'NA') not in ['NA', 'No']:
            user_selections.append(('Human Agents', inputs['human_agents']))
        if inputs.get('ai_module', 'NA') not in ['NA', 'No']:
            user_selections.append(('AI Module', inputs['ai_module']))
        if inputs.get('smart_cpaas', 'No') == 'Yes':
            user_selections.append(('Smart CPaaS', 'Yes'))
        if inputs.get('increased_tps', 'NA') not in ['NA', 'No']:
            user_selections.append(('Increased TPS', inputs['increased_tps']))
        inclusions = initialize_inclusions()
        final_inclusions = []
        # Always include base inclusions except those that overlap with selected options
        # Ensure these variables are defined before use
        personalize_load = inputs.get('personalize_load', 'NA')
        human_agents = inputs.get('human_agents', 'NA')
        increased_tps = inputs.get('increased_tps', 'NA')
        overlap_map = {
       'Personalize Load': ['Personalize Lite (upto 1ml and no advanced events)'],
       'Human Agents': ['Agent Assist <20'],
       'Increased TPS': ['80 TPS']
   }
        base_inclusions = inclusions['Platform Fee Used for Margin Calculation']
        # Remove overlapping items from base inclusions
        base_inclusions_filtered = []
        # Map of option to their possible overlapping base inclusions
        overlap_map = {
            'Personalize Load': ['Personalize Lite (upto 1ml and no advanced events)'],
            'Human Agents': ['Agent Assist <20'],
            'Increased TPS': ['80 TPS'],
        }
        # Get selected options
        personalize_load = inputs.get('personalize_load', 'NA')
        human_agents = inputs.get('human_agents', 'NA')
        increased_tps = inputs.get('increased_tps', 'NA')
        # Only add base inclusions that do not overlap with selected options
        for item in base_inclusions:
            if (personalize_load != 'Lite' and item in overlap_map['Personalize Load']) or \
               (human_agents not in ['NA', 'No', '<20'] and item in overlap_map['Human Agents']) or \
               (increased_tps not in ['NA', 'No', '80'] and item in overlap_map['Increased TPS']):
                continue
            base_inclusions_filtered.append(item)
        final_inclusions += base_inclusions_filtered
        # Add only the selected option inclusions
        if personalize_load == 'Advanced':
            final_inclusions += inclusions['Personalize Load Advanced']
        elif personalize_load == 'Standard':
            final_inclusions += inclusions['Personalize Load Standard']
        elif personalize_load == 'Lite':
            final_inclusions += inclusions['Personalize Load Lite']
        # Human Agents
        agent_inclusion_added = False
        if human_agents == '100+':
            final_inclusions += inclusions.get('Human Agents 100+', [])
            agent_inclusion_added = True
        elif human_agents == '50+':
            final_inclusions += inclusions.get('Human Agents 50+', [])
            agent_inclusion_added = True
        elif human_agents == '20+':
            final_inclusions += inclusions.get('Human Agents 20+', [])
            agent_inclusion_added = True
        elif human_agents == '<20':
            final_inclusions += inclusions.get('Human Agents <20', [])
            agent_inclusion_added = True
        # Fallback: if a Human Agents value is selected but not found, show generic inclusion
        if not agent_inclusion_added and human_agents not in ['NA', 'No', '']:
            final_inclusions.append('Agent Assist included')
        # Increased TPS
        if increased_tps == '1000':
            final_inclusions += inclusions['Increased TPS 1000']
        elif increased_tps == '250':
            final_inclusions += inclusions['Increased TPS 250']
        else:
            if '80 TPS' not in final_inclusions:
                final_inclusions.append('80 TPS')
        # AI Module
        if inputs.get('ai_module', 'No') == 'Yes':
            final_inclusions += inclusions['AI Module Yes']
            # If a specific AI agent model is selected, include it explicitly (without complexity text)
            ai_model = inputs.get('ai_agent_model', '')
            if ai_model and ai_model != 'None':
                final_inclusions.append(f"AI responses powered by {ai_model}")
        # Smart CPaaS
        if inputs.get('smart_cpaas', 'No') == 'Yes':
            final_inclusions += inclusions['Smart CPaaS Yes']
        # Voice Notes
        if inputs.get('voice_notes_price', 'No') == 'Yes':
            final_inclusions.append('Voice notes processing through STT/TTS features')
        # BFSI Tier (highest only)
        bfsi_tier = inputs.get('bfsi_tier', 'NA')
        if bfsi_tier == 'Tier 3':
            final_inclusions += inclusions['BFSI Tier 3']
        elif bfsi_tier == 'Tier 2':
            final_inclusions += inclusions['BFSI Tier 2']
        elif bfsi_tier == 'Tier 1':
            final_inclusions += inclusions['BFSI Tier 1']
        # Remove lower-tier inclusions if higher tier is selected
        if human_agents in ['20+', '50+', '100+']:
            if 'Agent Assist < 20 Agents' in final_inclusions:
                final_inclusions.remove('Agent Assist < 20 Agents')
        if personalize_load in ['Standard', 'Advanced']:
            if 'personalize lite upto 1 million records - no advanced events' in final_inclusions:
                final_inclusions.remove('personalize lite upto 1 million records - no advanced events')
        if increased_tps in ['250', '1000']:
            if '80 TPS' in final_inclusions:
                final_inclusions.remove('80 TPS')
        # Remove duplicates
        final_inclusions = list(dict.fromkeys(final_inclusions))
        # Pass contradiction_warning to the template for display if needed
        session['selected_components'] = user_selections
        session['results'] = results
        session['chosen_platform_fee'] = chosen_platform_fee
        session['rate_card_platform_fee'] = rate_card_platform_fee
        session['user_selections'] = user_selections
        session['inclusions'] = inclusions
        session['dev_cost_breakdown'] = dev_cost_breakdown
        session['dev_cost_currency'] = dev_cost_currency
        session['manday_breakdown'] = manday_breakdown
        session['manday_rates'] = manday_rates
        # Add platform fee as a line item if not already present (for volume path)
        platform_fee_line = {
            'line_item': 'Platform Fee (Chosen)',
            'volume': '',
            'chosen_price': '',
            'suggested_price': '',
            'overage_price': '',
            'revenue': float(platform_fee)
        }
        if not any(item.get('line_item') == 'Platform Fee (Chosen)' for item in results['line_items']):
            results['line_items'].append(platform_fee_line)
        # Create margin table for volume-based path
        margin_line_items = []
        # --- Determine rate card markups for India ---
        rate_card_markups = {}
        if inputs.get('country') == 'India':
            # Committed amount route
            committed_amount = float(inputs.get('committed_amount', 0) or 0)
            if committed_amount > 0:
                ca_rates = get_committed_amount_rates(inputs.get('country', 'India'), committed_amount)
                rate_card_markups = {
                    'AI Message': ca_rates['ai'],
                    'Advanced Message': ca_rates['advanced'],
                    'Basic Marketing Message': ca_rates['basic_marketing'],
                    'Basic Utility/Authentication Message': ca_rates['basic_utility'],
                }
        # ---
        for item in results['line_items']:
            if item.get('line_item') != 'Platform Fee (Chosen)':
                chosen_price = item.get('chosen_price', 0)
                # Use correct rate card markup for India
                rate_card_key = item.get('line_item') or item.get('label', '')
                if inputs.get('country') == 'India' and rate_card_key in rate_card_markups:
                    suggested_price = rate_card_markups[rate_card_key]
                else:
                    suggested_price = item.get('suggested_price', 0)
                discount_percent = ''
                if chosen_price and suggested_price and float(suggested_price) > 0:
                    try:
                        discount_percent = f"{((float(suggested_price) - float(chosen_price)) / float(suggested_price) * 100):.3f}%"
                    except Exception:
                        discount_percent = '0.00%'
                else:
                    discount_percent = '0.00%'
                margin_line_items.append({
                    'line_item': rate_card_key,
                    'chosen_price': chosen_price,
                    'rate_card_price': suggested_price,
                    'discount_percent': discount_percent
                })
        # Add platform fee to margin table with correct discount percent
        pf_discount = '0.00%'
        try:
            if rate_card_platform_fee and float(rate_card_platform_fee) > 0:
                pf_discount = f"{((float(rate_card_platform_fee) - float(platform_fee)) / float(rate_card_platform_fee) * 100):.3f}%"
        except Exception:
            pf_discount = '0.00%'
        margin_line_items.append({
            'line_item': 'Platform Fee',
            'chosen_price': platform_fee,
            'rate_card_price': rate_card_platform_fee,
            'discount_percent': pf_discount
        })
        # Helper to format numbers for display
        def fmt(val):
            # Format numbers: no decimals if .00, else show up to 4 decimals
            try:
                if isinstance(val, str) and val.replace('.', '', 1).isdigit():
                    val = float(val)
                if isinstance(val, (int, float)):
                    # Check if the number is a whole number
                    if val == int(val):
                        return str(int(val))
                    else:
                        # Show up to 4 decimal places, but remove trailing zeros
                        s = f"{val:.4f}"
                        return s.rstrip('0').rstrip('.')
            except Exception:
                pass
            return val
        # Format pricing_line_items
        for item in results['line_items']:
            for key in ['chosen_price', 'overage_price', 'suggested_price', 'revenue', 'volume']:
                if key in item and (isinstance(item[key], (int, float)) or (isinstance(item[key], str) and item[key].replace('.', '', 1).isdigit())) and item[key] != '':
                    item[key] = fmt(item[key])
        # Format margin_line_items
        for item in margin_line_items:
            for key in ['chosen_price', 'rate_card_price']:
                if key in item and (isinstance(item[key], (int, float)) or (isinstance(item[key], str) and item[key].replace('.', '', 1).isdigit())) and item[key] != '':
                    item[key] = fmt(item[key])
        # Log analytics for volumes flow
        # Prefer profile data where available, falling back to legacy inputs
        profile = session.get('profile') or {}
        analytics_kwargs = dict(
            timestamp=datetime.utcnow(),
            user_name=profile.get('name', inputs.get('user_name', '')),
            user_email=profile.get('email', None),
            country=profile.get('country', inputs.get('country', '')),
            region=profile.get('region', inputs.get('region', '')),
            platform_fee=platform_fee,
            ai_price=ai_price,
            advanced_price=advanced_price,
            basic_marketing_price=basic_marketing_price,
            basic_utility_price=basic_utility_price,
            currency=COUNTRY_CURRENCY.get(inputs.get('country', 'India'), '$'),
            ai_rate_card_price=suggested_ai,
            advanced_rate_card_price=suggested_advanced,
            basic_marketing_rate_card_price=suggested_marketing,
            basic_utility_rate_card_price=suggested_utility,
            ai_volume=ai_volume,
            advanced_volume=advanced_volume,
            basic_marketing_volume=basic_marketing_volume,
            basic_utility_volume=basic_utility_volume,
            # AI agent selection (model + complexity) for analytics
            ai_agent_model=inputs.get('ai_agent_model', ''),
            ai_agent_complexity=inputs.get('ai_agent_complexity', 'regular'),
            # Store user-submitted manday rates
            bot_ui_manday_rate=manday_rates.get('bot_ui'),
            custom_ai_manday_rate=manday_rates.get('custom_ai'),
            bot_ui_mandays=manday_breakdown.get('bot_ui', 0),
            custom_ai_mandays=manday_breakdown.get('custom_ai', 0),
            committed_amount=inputs.get('committed_amount', None),
            # Voice notes fields
            voice_notes_price=inputs.get('voice_notes_price', 'No'),
            voice_notes_model=inputs.get('voice_notes_model', ''),
            voice_notes_rate=float(inputs.get('voice_notes_rate', 0)) if inputs.get('voice_notes_rate') else None,
        )
        # Voice channel fields (if present)
        analytics_kwargs['channel_type'] = inputs.get('channel_type')
        try:
            vp = results.get('voice_pricing', None)
            if vp:
                analytics_kwargs.update({
                    'voice_mandays': vp.get('voice_mandays'),
                    'voice_dev_cost': vp.get('voice_dev_cost'),
                    'voice_platform_fee': vp.get('voice_platform_fee'),
                    'whatsapp_setup_fee': vp.get('whatsapp_setup_fee'),
                    'voice_cost_pstn_inbound': vp.get('calling_costs', {}).get('pstn_inbound_ai'),
                    'voice_cost_pstn_outbound': vp.get('calling_costs', {}).get('pstn_outbound_ai'),
                    'voice_cost_pstn_manual': vp.get('calling_costs', {}).get('pstn_manual_c2c'),
                    'voice_cost_wa_outbound': vp.get('calling_costs', {}).get('whatsapp_voice_outbound'),
                    'voice_cost_wa_inbound': vp.get('calling_costs', {}).get('whatsapp_voice_inbound'),
                    'voice_total_cost': vp.get('total_voice_cost'),
                })
                # Minutes from inputs
                for key in [
                    'pstn_inbound_ai_minutes','pstn_inbound_committed','pstn_outbound_ai_minutes','pstn_outbound_committed',
                    'pstn_manual_minutes','pstn_manual_committed','whatsapp_voice_outbound_minutes','whatsapp_voice_inbound_minutes']:
                    analytics_kwargs[key] = float(inputs.get(key, 0) or 0)
        except Exception as _:
            pass
        analytics_kwargs['calculation_id'] = session.get('calculation_id')
        # Set calculation_route for analytics
        if all(float(inputs.get(v, 0)) == 0.0 for v in ['ai_volume', 'advanced_volume', 'basic_marketing_volume', 'basic_utility_volume']) and float(inputs.get('committed_amount', 0) or 0) > 0:
            analytics_kwargs['calculation_route'] = 'bundle'
        else:
            analytics_kwargs['calculation_route'] = 'volumes'
        # Debug: Log manday rates and breakdown before saving to Analytics
        print(f"DEBUG: Saving Analytics: bot_ui_manday_rate={analytics_kwargs.get('bot_ui_manday_rate')}, custom_ai_manday_rate={analytics_kwargs.get('custom_ai_manday_rate')}, bot_ui_mandays={analytics_kwargs.get('bot_ui_mandays')}, custom_ai_mandays={analytics_kwargs.get('custom_ai_mandays')}, calculation_route={analytics_kwargs.get('calculation_route')}", file=sys.stderr, flush=True)
        new_analytics = Analytics(**analytics_kwargs)
        db.session.add(new_analytics)
        db.session.commit()
        # Top 5 users by number of calculations
        user_names = [row[0] for row in db.session.query(Analytics.user_name).all() if row[0]]
        top_users = Counter(user_names).most_common()  # Remove the 5 limit
        # When setting results['suggested_revenue'], use rate_card_platform_fee instead of platform_fee
        results['suggested_revenue'] = (results.get('suggested_revenue', 0) - platform_fee) + rate_card_platform_fee
        print("RENDERING RESULTS PAGE", file=sys.stderr, flush=True)
        contradiction_warning = None
        # Determine if SOW beta should be enabled for this user
        profile = session.get('profile') or {}
        email_for_sow = (profile.get('email') or '').strip().lower()
        sow_beta_enabled = bool(email_for_sow and email_for_sow in SOW_BETA_EMAILS)
        # Always calculate final_price_details for all routes
        committed_amount = float(inputs.get('committed_amount', 0) or 0)
        country = inputs.get('country', 'India')
        ai_volume = float(inputs.get('ai_volume', 0) or 0)
        advanced_volume = float(inputs.get('advanced_volume', 0) or 0)
        platform_fee_total = float(platform_fee)
        meta_costs = meta_costs_table.get(country, meta_costs_table['APAC'])
        if all(float(inputs.get(v, 0)) == 0.0 for v in ['ai_volume', 'advanced_volume', 'basic_marketing_volume', 'basic_utility_volume']):
            rates = get_committed_amount_rates(country, committed_amount)
            ai_price = rates['ai']
            advanced_price = rates['advanced']
        else:
            ai_price = session.get('pricing_inputs', {}).get('ai_price', 0)
            advanced_price = session.get('pricing_inputs', {}).get('advanced_price', 0)
        marketing_price = session.get('pricing_inputs', {}).get('basic_marketing_price', 0)
        utility_price = session.get('pricing_inputs', {}).get('basic_utility_price', 0)
        final_price_details = {
            'platform_fee_total': platform_fee_total,
            'fixed_platform_fee': platform_fee,
            'ai_messages': {
                'volume': int(ai_volume),
                'price_per_msg': round(ai_price + meta_costs['ai'], 4),
                'markup_per_msg': round(ai_price, 4),
                'overage_price': calculate_safe_overage_price(ai_price, meta_costs['ai'])
            },
            'advanced_messages': {
                'volume': int(advanced_volume),
                'price_per_msg': round(advanced_price + meta_costs.get('advanced', 0), 4),
                'markup_per_msg': round(advanced_price, 4),
                'overage_price': calculate_safe_overage_price(advanced_price, meta_costs.get('advanced', 0))
            },
            'marketing_message': {
                'volume': int(inputs.get('basic_marketing_volume', 0) or 0),
                'price_per_msg': round(marketing_price + meta_costs.get('marketing', 0), 4),
                'markup_per_msg': round(marketing_price, 4),
                'overage_price': calculate_safe_overage_price(marketing_price, meta_costs.get('marketing', 0))
            },
            'utility_message': {
                'volume': int(inputs.get('basic_utility_volume', 0) or 0),
                'price_per_msg': round(utility_price + meta_costs.get('utility', 0), 4),
                'markup_per_msg': round(utility_price, 4),
                'overage_price': calculate_safe_overage_price(utility_price, meta_costs.get('utility', 0))
            }
        }
        print("DEBUG: advanced_price =", advanced_price)
        print("DEBUG: meta_costs['advanced'] =", meta_costs.get('advanced', 0))
        print("DEBUG: final_price_details.advanced_messages.price_per_msg =", final_price_details['advanced_messages']['price_per_msg'])
        if all(float(inputs.get(v, 0)) == 0.0 for v in ['ai_volume', 'advanced_volume', 'basic_marketing_volume', 'basic_utility_volume']):
            committed_amount = float(inputs.get('committed_amount', 0) or 0)
            rates = get_committed_amount_rates(inputs.get('country', 'India'), committed_amount)
            meta_costs = meta_costs_table.get(inputs.get('country', 'India'), meta_costs_table['APAC'])
            ai_volume = float(inputs.get('ai_volume', 0) or 0)
            advanced_volume = float(inputs.get('advanced_volume', 0) or 0)
            basic_marketing_volume = float(inputs.get('basic_marketing_volume', 0) or 0)
            basic_utility_volume = float(inputs.get('basic_utility_volume', 0) or 0)
            platform_fee_total = float(platform_fee)  # You can adjust this if you want to sum other fees
            final_price_details = {
                'platform_fee_total': platform_fee_total,
                'fixed_platform_fee': platform_fee,
                'ai_messages': {
                    'volume': int(ai_volume),
                    'price_per_msg': round(rates['ai'] + meta_costs['ai'], 4),
                    'markup_per_msg': round(rates['ai'], 4),
                    'overage_price': calculate_safe_overage_price(rates['ai'], meta_costs['ai'])
                },
                'advanced_messages': {
                    'volume': int(advanced_volume),
                    'price_per_msg': round(rates['advanced'] + meta_costs.get('advanced', 0), 4),
                    'markup_per_msg': round(rates['advanced'], 4),
                    'overage_price': calculate_safe_overage_price(rates['advanced'], meta_costs.get('advanced', 0))
                },
                'marketing_message': {
                    'volume': int(inputs.get('basic_marketing_volume', 0) or 0),
                    'price_per_msg': round(session.get('pricing_inputs', {}).get('basic_marketing_price', 0) + meta_costs.get('marketing', 0), 4),
                    'markup_per_msg': round(session.get('pricing_inputs', {}).get('basic_marketing_price', 0), 4),
                    'overage_price': calculate_safe_overage_price(session.get('pricing_inputs', {}).get('basic_marketing_price', 0), meta_costs.get('marketing', 0))
                },
                'utility_message': {
                    'volume': int(inputs.get('basic_utility_volume', 0) or 0),
                    'price_per_msg': round(session.get('pricing_inputs', {}).get('basic_utility_price', 0) + meta_costs.get('utility', 0), 4),
                    'markup_per_msg': round(session.get('pricing_inputs', {}).get('basic_utility_price', 0), 4),
                    'overage_price': calculate_safe_overage_price(session.get('pricing_inputs', {}).get('basic_utility_price', 0), meta_costs.get('utility', 0))
                }
            }
            # Calculate pricing simulation for internal analysis
            pricing_simulation = calculate_pricing_simulation(inputs, session.get('pricing_inputs'))
            # Persist for SOW downloads and results refresh
            session['final_price_details'] = final_price_details
            session['final_inclusions'] = final_inclusions
            return render_template(
                'index.html',
                step='results',
                currency_symbol=currency_symbol,
                inclusions=final_inclusions,
                final_inclusions=final_inclusions,
                results=results,
                bundle_details=bundle_details,
                expected_invoice_amount=expected_invoice_amount,
                chosen_platform_fee=chosen_platform_fee,
                rate_card_platform_fee=rate_card_platform_fee,
                platform_fee=platform_fee,
                platform_fee_rate_card=rate_card_platform_fee,
                pricing_table=results['line_items'],
                margin_table=margin_line_items,
                user_selections=user_selections,
                inputs=inputs,
                contradiction_warning=contradiction_warning,
                top_users=top_users,
                total_mandays=total_mandays,
                total_dev_cost=total_dev_cost,
                dev_cost_currency=dev_cost_currency,
                manday_breakdown=manday_breakdown,
                manday_rates=manday_rates,
                calculation_id=calculation_id,
                dev_cost_breakdown=dev_cost_breakdown,
                final_price_details=final_price_details,
                pricing_simulation=pricing_simulation,
                sow_beta_enabled=sow_beta_enabled,
                platform_pricing_guidance=PLATFORM_PRICING_GUIDANCE,
                min_fees=min_fees
            )
        # Recalculate platform fee for the current country and selections before rendering results
        country = inputs.get('country', 'India')
        platform_fee, fee_currency = calculate_platform_fee(
            country,
            inputs.get('bfsi_tier', 'NA'),
            inputs.get('personalize_load', 'NA'),
            inputs.get('human_agents', 'NA'),
            inputs.get('ai_module', 'NA'),
            inputs.get('smart_cpaas', 'No'),
            inputs.get('increased_tps', 'NA')
        )
        session['chosen_platform_fee'] = platform_fee
        if 'pricing_inputs' in session:
            session['pricing_inputs']['platform_fee'] = platform_fee
        # Use this platform_fee for all downstream calculations and rendering
        pricing_simulation = calculate_pricing_simulation(inputs, session.get('pricing_inputs'))
        # Persist for SOW downloads and results refresh
        session['final_price_details'] = final_price_details
        session['final_inclusions'] = final_inclusions
        # Record successful results step for funnel analytics
        record_funnel_event('results', inputs=inputs, profile=session.get('profile') or {})
        return render_template(
            'index.html',
            step='results',
            currency_symbol=currency_symbol,
            inclusions=final_inclusions,
            final_inclusions=final_inclusions,
            results=results,
            bundle_details=bundle_details,
            expected_invoice_amount=expected_invoice_amount,
            chosen_platform_fee=chosen_platform_fee,
            rate_card_platform_fee=rate_card_platform_fee,
            platform_fee=chosen_platform_fee,
            platform_fee_rate_card=rate_card_platform_fee,
            pricing_table=results['line_items'],
            margin_table=margin_line_items,
            user_selections=user_selections,
            inputs=inputs,
            contradiction_warning=contradiction_warning,
            top_users=top_users,
            total_mandays=total_mandays,
            total_dev_cost=total_dev_cost,
            dev_cost_currency=dev_cost_currency,
            manday_breakdown=manday_breakdown,
            manday_rates=manday_rates,
            calculation_id=calculation_id,
            dev_cost_breakdown=dev_cost_breakdown,
            final_price_details=final_price_details,
            pricing_simulation=pricing_simulation,
            sow_beta_enabled=sow_beta_enabled,
            platform_pricing_guidance=PLATFORM_PRICING_GUIDANCE,
            min_fees=min_fees
        )
    # Defensive: handle GET or POST for edit actions
    elif step == 'volumes':
        inputs = session.get('inputs', {}) or {}
        profile = session.get('profile') or {}
        # If we have a profile but no inputs yet, prefill name/country/region from profile
        if profile and not inputs:
            if profile.get('name'):
                inputs['user_name'] = profile['name']
            if profile.get('email'):
                inputs['user_email'] = profile['email']
            if profile.get('country'):
                inputs['country'] = profile['country']
            if profile.get('region') is not None:
                inputs['region'] = profile['region']
            session['inputs'] = inputs
        if not inputs:
            if request.method == 'POST':
                flash('Session expired or missing. Please start again.', 'error')
            currency_symbol = COUNTRY_CURRENCY.get('India', '₹')
            record_funnel_event('volumes', inputs={}, profile=profile)
            return render_template('index.html', step='volumes', currency_symbol=currency_symbol, inputs={}, profile=profile, calculation_id=calculation_id, min_fees=min_fees, ai_agent_pricing=AI_AGENT_PRICING)
        currency_symbol = COUNTRY_CURRENCY.get(inputs.get('country', 'India'), '$')
        record_funnel_event('volumes', inputs=inputs, profile=profile)
        return render_template('index.html', step='volumes', currency_symbol=currency_symbol, inputs=inputs, profile=profile, calculation_id=calculation_id, min_fees=min_fees, ai_agent_pricing=AI_AGENT_PRICING)
    elif step == 'prices':
        inputs = session.get('inputs', {})
        pricing_inputs = session.get('pricing_inputs', {}) or {}
        country = inputs.get('country', 'India')
        dev_location = inputs.get('dev_location', 'India')
        # Get default rates
        country = inputs.get('country', 'India').strip()
        print(f"DEBUG: country used for manday rates = '{country}'", file=sys.stderr, flush=True)
        rates = COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])
        if country == 'LATAM':
            default_bot_ui = float(rates['bot_ui'].get(dev_location, 0.0) or 0.0)
            default_custom_ai = float(rates['custom_ai'].get(dev_location, 0.0) or 0.0)
        else:
            default_bot_ui = float(rates.get('bot_ui', 0.0) or 0.0)
            default_custom_ai = float(rates.get('custom_ai', 0.0) or 0.0)
        if request.method == 'POST':
            # Validate user rates
            def parse_number(val):
                try:
                    return float(str(val).replace(',', '')) if val and str(val).strip() else 0.0
                except Exception:
                    return 0.0
            user_bot_ui = parse_number(request.form.get('bot_ui_manday_rate', default_bot_ui))
            user_custom_ai = parse_number(request.form.get('custom_ai_manday_rate', default_custom_ai))
            if user_bot_ui < 0.5 * default_bot_ui or user_custom_ai < 0.5 * default_custom_ai:
                flash('Manday rates cannot be discounted by more than 50% from the default rate.', 'error')
                record_funnel_event('prices', inputs=inputs)
                return render_template('index.html', step='prices', suggested={
                    'ai_price': pricing_inputs.get('ai_price', ''),
                    'advanced_price': pricing_inputs.get('advanced_price', ''),
                    'basic_marketing_price': pricing_inputs.get('basic_marketing_price', ''),
                    'basic_utility_price': pricing_inputs.get('basic_utility_price', ''),
                    'bot_ui_manday_rate': default_bot_ui,
                    'custom_ai_manday_rate': default_custom_ai,
                }, inputs=inputs, currency_symbol=COUNTRY_CURRENCY.get(country, '$'), platform_fee=pricing_inputs.get('platform_fee', inputs.get('platform_fee', '')), calculation_id=calculation_id, min_fees=min_fees, ai_agent_pricing=AI_AGENT_PRICING)
            # Save user rates for use in results
            session['manday_rates'] = {
                'bot_ui': user_bot_ui,
                'custom_ai': user_custom_ai,
                'default_bot_ui': default_bot_ui,
                'default_custom_ai': default_custom_ai,
            }
        else:
            # GET: pre-fill with defaults
            record_funnel_event('prices', inputs=inputs)
            return render_template('index.html', step='prices', suggested={
                'ai_price': pricing_inputs.get('ai_price', ''),
                'advanced_price': pricing_inputs.get('advanced_price', ''),
                'basic_marketing_price': pricing_inputs.get('basic_marketing_price', ''),
                'basic_utility_price': pricing_inputs.get('basic_utility_price', ''),
                'bot_ui_manday_rate': default_bot_ui,
                'custom_ai_manday_rate': default_custom_ai,
            }, inputs=inputs, currency_symbol=COUNTRY_CURRENCY.get(country, '$'), platform_fee=pricing_inputs.get('platform_fee', inputs.get('platform_fee', '')), calculation_id=calculation_id, min_fees=min_fees)
    elif step == 'bundle' and request.method == 'POST':
        # User submitted messaging bundle commitment (can be 0)
        inputs = session.get('inputs', {}) or {}
        committed_amount = request.form.get('committed_amount', '0')
        # Remove commas if present
        try:
            committed_amount = float(str(committed_amount).replace(',', ''))
        except Exception:
            committed_amount = 0.0
        inputs['committed_amount'] = committed_amount
        session['inputs'] = inputs
        # Recalculate platform fee for the current country and selections
        country = inputs.get('country', 'India')
        platform_fee, fee_currency = calculate_platform_fee(
            country,
            inputs.get('bfsi_tier', 'NA'),
            inputs.get('personalize_load', 'NA'),
            inputs.get('human_agents', 'NA'),
            inputs.get('ai_module', 'NA'),
            inputs.get('smart_cpaas', 'No'),
            inputs.get('increased_tps', 'NA')
        )
        session['chosen_platform_fee'] = platform_fee
        # Go to prices page
        pricing_inputs = session.get('pricing_inputs', {}) or {}
        dev_location = inputs.get('dev_location', 'India')
        rates = COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])
        if country == 'LATAM':
            default_bot_ui = float(rates['bot_ui'].get(dev_location, 0.0) or 0.0)
            default_custom_ai = float(rates['custom_ai'].get(dev_location, 0.0) or 0.0)
        else:
            default_bot_ui = float(rates.get('bot_ui', 0.0) or 0.0)
            default_custom_ai = float(rates.get('custom_ai', 0.0) or 0.0)
        # --- Set default per-message prices for all countries based on committed amount using committed_amount_slabs ---
        slabs = committed_amount_slabs.get(country, committed_amount_slabs.get('APAC', []))
        # Find the correct slab for the committed amount
        selected_slab = None
        for lower, upper, rates in slabs:
            if lower <= committed_amount < upper:
                selected_slab = rates
                break
        if not selected_slab and slabs:
            # If above all slabs, use the highest slab
            selected_slab = slabs[-1][2]
        suggested_prices = {
            'ai_price': selected_slab['ai'] if selected_slab else '',
            'advanced_price': selected_slab['advanced'] if selected_slab else '',
            'basic_marketing_price': selected_slab['basic_marketing'] if selected_slab else '',
            'basic_utility_price': selected_slab['basic_utility'] if selected_slab else '',
            'bot_ui_manday_rate': default_bot_ui,
            'custom_ai_manday_rate': default_custom_ai,
        }
        suggested_prices = patch_suggested_prices(suggested_prices, inputs)
        record_funnel_event('prices', inputs=inputs)
        return render_template('index.html', step='prices', suggested=suggested_prices, inputs=inputs, currency_symbol=COUNTRY_CURRENCY.get(country, '$'), platform_fee=pricing_inputs.get('platform_fee', inputs.get('platform_fee', '')), calculation_id=calculation_id, min_fees=min_fees)
    elif step == 'results':
        # Handle GET request for results page (page refresh)
        inputs = session.get('inputs', {})
        if not inputs or not session.get('results'):
            flash('Session expired or missing. Please start again.', 'error')
            currency_symbol = COUNTRY_CURRENCY.get('India', '₹')
            record_funnel_event('volumes', inputs={}, profile=session.get('profile') or {})
            return render_template('index.html', step='volumes', currency_symbol=currency_symbol, inputs={}, calculation_id=calculation_id, min_fees=min_fees)
        
        # Re-render results page with existing session data
        country = inputs.get('country', 'India')
        currency_symbol = COUNTRY_CURRENCY.get(country, '$')
        
        # Get all the data from session
        results = session.get('results', {})
        pricing_inputs = session.get('pricing_inputs', {})
        final_inclusions = session.get('inclusions', {})
        dev_cost_breakdown = session.get('dev_cost_breakdown', {})
        final_price_details = session.get('final_price_details', {})
        manday_breakdown = session.get('manday_breakdown', {})
        manday_rates = session.get('manday_rates', {})
        dev_cost_currency = session.get('dev_cost_currency', 'INR')
        
        # If dev_cost_breakdown is missing or empty, calculate it
        if not dev_cost_breakdown or 'total_cost' not in dev_cost_breakdown:
            try:
                total_dev_cost, dev_cost_currency, dev_cost_breakdown = calculate_total_manday_cost(inputs, manday_rates)
                manday_breakdown = dev_cost_breakdown.get('mandays_breakdown', {})
                # Store in session for future use
                session['dev_cost_breakdown'] = dev_cost_breakdown
                session['dev_cost_currency'] = dev_cost_currency
                session['manday_breakdown'] = manday_breakdown
            except Exception as e:
                print(f"Error calculating dev_cost_breakdown: {e}", file=sys.stderr, flush=True)
                # Fallback to empty structure
                dev_cost_breakdown = {'total_cost': 0, 'ba_cost': 0, 'qa_cost': 0, 'pm_cost': 0, 'uplift_amount': 0}
                dev_cost_currency = 'INR'
        
        # Recalculate pricing simulation
        pricing_simulation = calculate_pricing_simulation(inputs, pricing_inputs)
        
        # Calculate total mandays for template
        total_mandays = calculate_total_mandays(inputs)
        
        # Ensure manday_breakdown has the correct structure
        if not manday_breakdown or 'bot_ui' not in manday_breakdown:
            manday_breakdown = calculate_total_mandays_breakdown(inputs)
        # Determine if SOW beta should be enabled for this user (refresh case)
        profile = session.get('profile') or {}
        email_for_sow = (profile.get('email') or '').strip().lower()
        sow_beta_enabled = bool(email_for_sow and email_for_sow in SOW_BETA_EMAILS)
        record_funnel_event('results', inputs=inputs, profile=profile)
        return render_template(
            'index.html',
            step='results',
            currency_symbol=currency_symbol,
            inclusions=final_inclusions,
            final_inclusions=final_inclusions,
            results=results,
            inputs=inputs,
            pricing_inputs=pricing_inputs,
            dev_cost_breakdown=dev_cost_breakdown,
            final_price_details=final_price_details,
            manday_breakdown=manday_breakdown,
            manday_rates=manday_rates,
            dev_cost_currency=dev_cost_currency,
            calculation_id=calculation_id,
            pricing_simulation=pricing_simulation,
            platform_pricing_guidance=PLATFORM_PRICING_GUIDANCE,
            min_fees=min_fees,
            total_mandays=total_mandays,
            sow_beta_enabled=sow_beta_enabled
        )
    # Default: show volume input form (this handles "Start Over" and initial page load)
    # Clear calculation_id when starting fresh
    if 'calculation_id' in session:
        session.pop('calculation_id')
        print("DEBUG: Cleared calculation_id for fresh start", file=sys.stderr, flush=True)
    
    country = session.get('inputs', {}).get('country', 'India')
    currency_symbol = COUNTRY_CURRENCY.get(country, '$')
    return render_template('index.html', step='volumes', currency_symbol=currency_symbol, inputs=session.get('inputs', {}), calculation_id=None, min_fees=min_fees, ai_agent_pricing=AI_AGENT_PRICING)


def generate_sow_docx(inputs, results, final_price_details, profile, sow_details=None, calculation_id=None):
    """
    Construct a Scope of Work .docx document in memory using the master template.
    """
    from pathlib import Path
    from datetime import datetime
    # Load master SOW template so all static sections stay intact
    template_path = Path(__file__).resolve().parent / 'sow_templates' / 'Latest SOW Master Copy August 2025.docx'
    doc = Document(str(template_path))

    # Basic profile / deal context (used where relevant)
    name = (profile or {}).get('name') or inputs.get('user_name', '')
    email = (profile or {}).get('email', '')
    country = (profile or {}).get('country') or inputs.get('country', '')
    region = (profile or {}).get('region') or inputs.get('region', '')

    # Update cover page fields if present (prepend simple prepared-for block near top)
    try:
        # Insert a new paragraph at the top with Prepared For
        first_para = doc.paragraphs[0]
        p = first_para.insert_paragraph_before()
    except Exception:
        p = doc.add_paragraph()
    run = p.add_run('Prepared For: ')
    run.bold = True
    p.add_run(name or 'Client')
    if email:
        p.add_run(f"  <{email}>")
    # Insert location and calculation ID just after
    if country or region:
        loc_line = ', '.join([v for v in [country, region] if v])
        doc.add_paragraph(f"Location: {loc_line}")
    if calculation_id:
        doc.add_paragraph(f"Calculation ID: {calculation_id}")

    # Revision History – fill first data row (Table 0)
    try:
        rev_table = doc.tables[0]
        if len(rev_table.rows) >= 2:
            rev_row = rev_table.rows[1]
            rev_row.cells[0].text = '1.0'
            rev_row.cells[1].text = datetime.utcnow().strftime('%Y-%m-%d')
            rev_row.cells[2].text = '1st draft'
            rev_row.cells[3].text = name or ''
            # Approver Name left blank
    except Exception:
        pass

    # Business Objectives Discovered – fill Table A in Section 1
    sow_details = sow_details or {}
    try:
        bo_table = doc.tables[3]  # Table 3: Business Objectives (Description / Client requirement)
        for row in bo_table.rows[1:]:
            desc = row.cells[0].text.strip()
            if desc == 'Business Objective':
                row.cells[1].text = sow_details.get('business_objective', '')
            elif desc == 'Business Problem':
                row.cells[1].text = sow_details.get('business_problem', '')
            elif desc == 'Use Case Narrative':
                row.cells[1].text = sow_details.get('use_case_narrative', '')
            elif desc == 'Expected Volumes':
                row.cells[1].text = sow_details.get('expected_volumes', '')
            elif desc == 'Target Audience':
                row.cells[1].text = sow_details.get('target_audience', '')
            elif desc.startswith('KPI'):
                row.cells[1].text = sow_details.get('kpis', '')
    except Exception:
        pass

    # Overall Requirement – Table A in Technical Requirements (Table 4)
    try:
        overall = doc.tables[4]
        channels_selected = sow_details.get('channels', []) or []
        # Helper to map channels into template groups
        def group_channels(group_names):
            return ', '.join([c for c in channels_selected if c in group_names])
        for row in overall.rows[1:]:
            desc = row.cells[0].text.strip()
            if desc == 'Number of journeys':
                val = sow_details.get('num_journeys') or inputs.get('num_journeys_price') or ''
                row.cells[1].text = str(val)
            elif desc == 'Channels':
                # Three buckets: WA/WA Voice, SMS/Others, Instagram/PSTN Voice
                row.cells[1].text = group_channels(['WhatsApp', 'WA', 'WhatsApp Voice', 'WA Voice'])
                row.cells[2].text = group_channels(['SMS', 'Others'])
                row.cells[3].text = group_channels(['Instagram', 'PSTN Voice'])
            elif desc == 'Bot Language':
                row.cells[1].text = sow_details.get('bot_language', '')
            elif desc.startswith('No. of APIs/Backend Services'):
                val = sow_details.get('num_apis') or inputs.get('num_apis_price') or ''
                row.cells[1].text = str(val)
            elif desc.startswith('Name of the backend systems in scope'):
                row.cells[1].text = sow_details.get('backend_systems', '')
            elif desc.startswith('Development Timelines'):
                val = sow_details.get('total_mandays', '')
                row.cells[1].text = str(val)
    except Exception:
        pass

    # AI Specifications – Table B (Table 5)
    try:
        ai_table = doc.tables[5]
        ai_module = inputs.get('ai_module', 'NA')
        if ai_module != 'Yes':
            # Omit AI Specifications table for non-AI SOWs
            tbl = ai_table._tbl
            tbl.getparent().remove(tbl)
        else:
            complexity = inputs.get('ai_agent_complexity', '') or sow_details.get('bot_complexity', '')
            model = inputs.get('ai_agent_model', '') or sow_details.get('llm_model', '')
            for row in ai_table.rows[1:]:
                desc = row.cells[0].text.strip()
                if desc == 'Bot Complexity':
                    row.cells[1].text = complexity
                elif desc == 'LLM Model':
                    row.cells[1].text = model
    except Exception:
        pass

    # Deployment / Hosting – Table D (Table 7)
    try:
        dep_table = doc.tables[7]
        dep_use_default = (sow_details.get('dep_use_default', 'yes') == 'yes')
        if dep_use_default:
            # Set default client requirement
            for row in dep_table.rows[1:]:
                desc = row.cells[0].text.strip()
                if desc == 'Deployment/Hosting':
                    row.cells[1].text = 'Gupshup India AWS'
                    break
            # Remove Questions and Presales Findings columns (last two)
            for col_idx in [3, 2]:
                for row in dep_table.rows:
                    cell = row.cells[col_idx]
                    cell._tc.getparent().remove(cell._tc)
    except Exception:
        pass

    # Scope / inclusions – we keep the static prose from the template, but can
    # optionally append dynamic inclusions list at the end of that section.
    inclusions = session.get('inclusions') or {}
    final_inclusions = session.get('final_inclusions') or []
    # Prefer the flattened list passed to the template; fall back to dict if needed.
    inclusion_items = final_inclusions or []
    if not inclusion_items and isinstance(inclusions, dict):
        for vals in inclusions.values():
            inclusion_items.extend(vals or [])
    if inclusion_items:
        # Append to end of document as an “Included Features” list so we
        # don’t disturb the master copy layout. Use default paragraph style
        # to avoid relying on template-specific bullet styles.
        doc.add_paragraph()  # spacer
        heading_par = doc.add_paragraph()
        run = heading_par.add_run('Included Features (from calculator):')
        run.bold = True
        for item in inclusion_items:
            doc.add_paragraph(str(item))

    # Capacity & Compliance – Table E (Table 8)
    try:
        cap = doc.tables[8]
        tps_default = (sow_details.get('tps_use_default', 'yes') == 'yes')
        dr_default = (sow_details.get('dr_use_default', 'yes') == 'yes')
        include_personalize = (sow_details.get('include_personalize', 'yes') == 'yes')
        additional_security = (sow_details.get('additional_security', 'no') == 'yes')

        rows_to_remove = []
        in_tps_block = False
        in_dr_block = False
        for idx, row in enumerate(cap.rows[1:], start=1):
            desc = row.cells[0].text.strip()
            # Track which block we're in (TPS/Data Retention) for multi-row sections
            if desc == 'TPS':
                in_tps_block = True
                in_dr_block = False
            elif desc == 'Data Retention':
                in_dr_block = True
                in_tps_block = False
            elif desc:
                in_tps_block = False
                in_dr_block = False

            if desc == 'TPS' and tps_default:
                row.cells[1].text = 'Standard 80'
                # Clear Questions and Presales columns as per instructions
                row.cells[2].text = ''
                row.cells[3].text = ''
            elif in_tps_block and tps_default:
                # For subsequent TPS rows, clear Questions/Presales content
                row.cells[2].text = ''
                row.cells[3].text = ''
            elif desc == 'Data Retention' and dr_default:
                row.cells[1].text = '2 years - Standard'
                row.cells[2].text = ''
                row.cells[3].text = ''
            elif in_dr_block and dr_default:
                # Subsequent Data Retention rows: clear Questions/Presales only
                row.cells[2].text = ''
                row.cells[3].text = ''

            if desc == 'Profile Storage Requirements' and not include_personalize:
                rows_to_remove.append(idx)
            elif desc == 'Security and Privacy Package' and not additional_security:
                rows_to_remove.append(idx)

        # Remove marked rows from bottom up
        for idx in sorted(rows_to_remove, reverse=True):
            cap._tbl.remove(cap.rows[idx]._tr)

        # If standard configs are used for both TPS and Data Retention,
        # remove the last two columns (Questions, Presales Findings) for
        # the entire table to simplify the layout.
        if tps_default and dr_default:
            for col_idx in [3, 2]:
                for row in cap.rows:
                    cell = row.cells[col_idx]
                    cell._tc.getparent().remove(cell._tc)
    except Exception:
        pass

    # Delivery – insert Presales name into credentials column (Table 11)
    try:
        roles = doc.tables[11]
        for row in roles.rows[1:]:
            if row.cells[0].text.strip() == 'Presales':
                row.cells[2].text = name or ''
                break
    except Exception:
        pass

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    try:
        logger.info(
            "SOW_GENERATED",
            extra={
                "calculation_id": calculation_id,
                "user_email": (profile or {}).get("email"),
                "country": (profile or {}).get("country") or inputs.get("country"),
            },
        )
    except Exception:
        # Logging must never break SOW generation
        pass
    return bio


@app.route('/generate-sow', methods=['GET'])
def generate_sow():
    """
    Generate and return a Scope of Work .docx based on the last completed calculation.
    """
    if not session.get('authenticated'):
        return redirect(url_for('login'))

    # Enforce SOW beta allowlist
    profile = session.get('profile') or {}
    email_for_sow = (profile.get('email') or '').strip().lower()
    if not email_for_sow or email_for_sow not in SOW_BETA_EMAILS:
        flash('SOW generation (beta) is currently enabled only for a limited set of internal users.', 'error')
        return redirect(url_for('index', step='results'))

    inputs = session.get('inputs')
    results = session.get('results')
    final_price_details = session.get('final_price_details')
    sow_details = session.get('sow_details') or {}
    profile = session.get('profile') or {}
    calculation_id = session.get('calculation_id')

    if not inputs or not results or not final_price_details:
        flash('No completed calculation found for SOW generation. Please run a pricing calculation first.', 'error')
        return redirect(url_for('index'))

    # Mark that the SOW was actually downloaded for this calculation
    if calculation_id:
        try:
            analytics_row = (
                Analytics.query.filter_by(calculation_id=calculation_id)
                .order_by(Analytics.timestamp.desc())
                .first()
            )
            if analytics_row:
                # Ensure click is marked as well for robustness
                if not analytics_row.sow_generate_clicked:
                    analytics_row.sow_generate_clicked = True
                analytics_row.sow_downloaded = True
                db.session.commit()
        except Exception:
            logger.exception("Failed to mark sow_downloaded in Analytics")

    try:
        logger.info(
            "SOW_GENERATE_REQUEST",
            extra={
                "calculation_id": calculation_id,
                "user_email": profile.get("email"),
                "country": inputs.get("country"),
                "route": "bundle" if float(inputs.get("committed_amount", 0) or 0) > 0 else "volumes",
            },
        )
    except Exception:
        pass

    # Record SOW download step for funnel analytics
    record_funnel_event('sow_download', inputs=inputs, profile=profile)

    docx_io = generate_sow_docx(inputs, results, final_price_details, profile, sow_details, calculation_id)
    filename = f"SOW_{calculation_id or 'pricing'}.docx"
    return send_file(
        docx_io,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route('/sow-details', methods=['GET', 'POST'])
def sow_details():
    """Intermediate SOW details page shown before generating the SOW doc."""
    if not session.get('authenticated'):
        return redirect(url_for('login'))

    profile = session.get('profile') or {}
    email_for_sow = (profile.get('email') or '').strip().lower()
    if not email_for_sow or email_for_sow not in SOW_BETA_EMAILS:
        flash('SOW generation (beta) is currently enabled only for a limited set of internal users.', 'error')
        return redirect(url_for('index', step='results'))

    inputs = session.get('inputs') or {}
    results = session.get('results')
    calculation_id = session.get('calculation_id')
    if not inputs or not results:
        flash('No completed calculation found for SOW generation. Please run a pricing calculation first.', 'error')
        return redirect(url_for('index'))

    # Record entry into SOW details step for funnel analytics
    record_funnel_event('sow_details', inputs=inputs, profile=profile)

    # Mark that the user has entered the SOW funnel for this calculation (for abandon analysis)
    if calculation_id:
        try:
            analytics_row = (
                Analytics.query.filter_by(calculation_id=calculation_id)
                .order_by(Analytics.timestamp.desc())
                .first()
            )
            if analytics_row and not analytics_row.sow_generate_clicked:
                analytics_row.sow_generate_clicked = True
                db.session.commit()
        except Exception:
            # Never break user flow if analytics update fails
            logger.exception("Failed to mark sow_generate_clicked in Analytics")

    existing = session.get('sow_details') or {}

    if request.method == 'POST':
        # Collect Business Objectives section
        sow = {
            'business_objective': request.form.get('business_objective', '').strip(),
            'business_problem': request.form.get('business_problem', '').strip(),
            'use_case_narrative': request.form.get('use_case_narrative', '').strip(),
            'expected_volumes': request.form.get('expected_volumes', '').strip(),
            'target_audience': request.form.get('target_audience', '').strip(),
            'kpis': request.form.get('kpis', '').strip(),
            # Overall technical numbers
            'num_journeys': request.form.get('num_journeys', '').strip(),
            'num_apis': request.form.get('num_apis', '').strip(),
            'bot_language': request.form.get('bot_language', '').strip(),
            'backend_systems': request.form.get('backend_systems', '').strip(),
            'total_mandays': request.form.get('total_mandays', '').strip(),
            # Channels multi-select
            'channels': request.form.getlist('channels'),
            # Deployment & capacity toggles
            'dep_use_default': request.form.get('dep_use_default', 'yes'),
            'tps_use_default': request.form.get('tps_use_default', 'yes'),
            'dr_use_default': request.form.get('dr_use_default', 'yes'),
            'include_personalize': request.form.get('include_personalize', 'yes'),
            'additional_security': request.form.get('additional_security', 'no'),
        }
        session['sow_details'] = sow
        try:
            logger.info(
                "SOW_DETAILS_SAVED",
                extra={
                    "user_email": profile.get("email"),
                    "country": inputs.get("country", ""),
                    "has_ai": inputs.get("ai_module") == "Yes",
                    "route": "bundle" if float(inputs.get("committed_amount", 0) or 0) > 0 else "volumes",
                },
            )
        except Exception:
            pass
        # After capturing details, trigger SOW generation
        return redirect(url_for('generate_sow'))

    # Defaults for initial load
    try:
        from calculator import calculate_total_mandays
        default_total_mandays = calculate_total_mandays(inputs)
    except Exception:
        default_total_mandays = ''
    defaults = {
        'num_journeys': inputs.get('num_journeys_price', ''),
        'num_apis': inputs.get('num_apis_price', ''),
        'total_mandays': default_total_mandays,
    }

    class Obj:  # simple helper to allow dot-access in template
        def __init__(self, d):
            for k, v in d.items():
                setattr(self, k, v)

    sow_obj = Obj(existing)
    defaults_obj = Obj(defaults)
    return render_template('sow_details.html', sow=sow_obj, defaults=defaults_obj)

@app.route('/analytics', methods=['GET', 'POST'])
def analytics():
    try:
        if request.method == 'POST':
            keyword = request.form.get('keyword', '')
            if keyword == SECRET_ANALYTICS_KEYWORD:
                # Query the database for analytics
                total_calculations = Analytics.query.count()
                # Calculations by day
                calculations_by_day = {str(row[0]): row[1] for row in db.session.query(func.date(Analytics.timestamp), func.count()).group_by(func.date(Analytics.timestamp)).all()}
                # Calculations by week
                calculations_by_week = {str(row[0]): row[1] for row in db.session.query(
                    func.to_char(Analytics.timestamp, 'IYYY-"W"IW'),
                    func.count()
                ).group_by(
                    func.to_char(Analytics.timestamp, 'IYYY-"W"IW')
                ).all()}
                # Most common countries
                from collections import Counter
                country_list = [row[0] for row in db.session.query(Analytics.country).all() if row[0] is not None]
                country_counter = Counter(country_list).most_common(5)
                # Platform fee stats
                platform_fees = [row[0] for row in db.session.query(Analytics.platform_fee).all() if row[0] is not None]
                if platform_fees:
                    avg_platform_fee = sum(platform_fees) / len(platform_fees)
                    min_platform_fee = min(platform_fees)
                    max_platform_fee = max(platform_fees)
                    median_platform_fee = sorted(platform_fees)[len(platform_fees)//2]
                else:
                    avg_platform_fee = min_platform_fee = max_platform_fee = median_platform_fee = 0
                # Platform fee options (most common platform_fee values)
                platform_fee_options = Counter(platform_fees).most_common(5)
                # Message price stats by type
                def get_stats(field):
                    vals = [getattr(a, field) for a in Analytics.query.all() if getattr(a, field) is not None]
                    if vals:
                        return {
                            'avg': sum(vals)/len(vals),
                            'min': min(vals),
                            'max': max(vals),
                            'median': sorted(vals)[len(vals)//2]
                        }
                    else:
                        return {'avg': 0, 'min': 0, 'max': 0, 'median': 0}
                ai_stats = get_stats('ai_price')
                advanced_stats = get_stats('advanced_price')
                marketing_stats = get_stats('basic_marketing_price')
                utility_stats = get_stats('basic_utility_price')
                # Voice notes stats
                voice_notes_stats = get_stats('voice_notes_rate')
                # Voice notes usage stats
                voice_notes_usage = Analytics.query.filter(Analytics.voice_notes_price == 'Yes').count()
                voice_notes_total = Analytics.query.count()
                voice_notes_percentage = (voice_notes_usage / voice_notes_total * 100) if voice_notes_total > 0 else 0
                # Message volume distribution (sum, min, max, median for each type)
                def get_volume_stats(field):
                    vals = [getattr(a, field) for a in Analytics.query.all() if getattr(a, field) is not None]
                    if vals:
                        return [sum(vals), min(vals), max(vals), sorted(vals)[len(vals)//2]]
                    else:
                        return [0, 0, 0, 0]
                message_volumes = {
                    'ai': get_volume_stats('ai_price'),
                    'advanced': get_volume_stats('advanced_price'),
                    'basic_marketing': get_volume_stats('basic_marketing_price'),
                    'basic_utility': get_volume_stats('basic_utility_price')
                }
                # Platform fee discount triggered (count)
                platform_fee_discount_triggered = db.session.query(Analytics).filter(Analytics.platform_fee < 0.3 * avg_platform_fee).count() if avg_platform_fee else 0
                # Margin chosen and rate card (dummy values for now, replace with real if available)
                margin_chosen = [95.825]  # Example value, replace with real calculation if available
                margin_rate_card = [0.0]  # Example value, replace with real calculation if available
                # Per-country stats for table
                stats = {}
                stats_by_region = {}
                countries = set(country_list)
                for country in countries:
                    country_analytics = Analytics.query.filter_by(country=country).all()
                    # Group by region within country
                    region_map = {}
                    for a in country_analytics:
                        region = getattr(a, 'region', '') or 'All'
                        if region == 'All' or not region:
                            region = 'All'  # Keep as 'All' instead of changing to country name
                        if region not in region_map:
                            region_map[region] = []
                        region_map[region].append(a)
                    stats_by_region[country] = {}
                    for region, region_analytics in region_map.items():
                        country_fees = [a.platform_fee for a in region_analytics if a.platform_fee is not None]
                        country_ai = [a.ai_price for a in region_analytics if a.ai_price is not None]
                        country_adv = [a.advanced_price for a in region_analytics if a.advanced_price is not None]
                        country_mark = [a.basic_marketing_price for a in region_analytics if a.basic_marketing_price is not None]
                        country_util = [a.basic_utility_price for a in region_analytics if a.basic_utility_price is not None]
                        country_voice_notes = [a.voice_notes_rate for a in region_analytics if a.voice_notes_rate is not None]
                        country_committed = [a.committed_amount for a in region_analytics if a.committed_amount not in (None, 0, '0', '', 'None')]
                        # --- One Time Dev Cost Aggregation ---
                        dev_costs = []
                        bot_ui_rates = [a.bot_ui_manday_rate for a in region_analytics if a.bot_ui_manday_rate not in (None, 0, '0', '', 'None')]
                        custom_ai_rates = [a.custom_ai_manday_rate for a in region_analytics if a.custom_ai_manday_rate not in (None, 0, '0', '', 'None')]
                        bot_ui_mandays = [getattr(a, 'bot_ui_mandays', None) for a in region_analytics]
                        custom_ai_mandays = [getattr(a, 'custom_ai_mandays', None) for a in region_analytics]
                        for a in region_analytics:
                            bot_ui_rate = a.bot_ui_manday_rate if a.bot_ui_manday_rate is not None else 0
                            custom_ai_rate = a.custom_ai_manday_rate if a.custom_ai_manday_rate is not None else 0
                            bot_days = getattr(a, 'bot_ui_mandays', 0) or 0
                            ai_days = getattr(a, 'custom_ai_mandays', 0) or 0
                            dev_cost = (bot_ui_rate * bot_days) + (custom_ai_rate * ai_days)
                            dev_costs.append(dev_cost)
                        def stat_dict(vals):
                            filtered = [float(v) for v in vals if v not in (None, 0, '0', '', 'None')]
                            if not filtered:
                                return {'avg': 0, 'min': 0, 'max': 0, 'median': 0}
                            avg = sum(filtered) / len(filtered)
                            min_v = min(filtered)
                            max_v = max(filtered)
                            sorted_vals = sorted(filtered)
                            n = len(sorted_vals)
                            if n % 2 == 1:
                                median = sorted_vals[n // 2]
                            else:
                                median = (sorted_vals[n // 2 - 1] + sorted_vals[n // 2]) / 2
                            return {'avg': avg, 'min': min_v, 'max': max_v, 'median': median}
                        stats_by_region[country][region] = {
                            'platform_fee': stat_dict(country_fees),
                            'msg_types': {
                                'ai': {
                                    'avg': sum(country_ai)/len(country_ai) if country_ai else 0,
                                    'min': min(country_ai) if country_ai else 0,
                                    'max': max(country_ai) if country_ai else 0,
                                    'median': sorted(country_ai)[len(country_ai)//2] if country_ai else 0
                                },
                                'advanced': {
                                    'avg': sum(country_adv)/len(country_adv) if country_adv else 0,
                                    'min': min(country_adv) if country_adv else 0,
                                    'max': max(country_adv) if country_adv else 0,
                                    'median': sorted(country_adv)[len(country_adv)//2] if country_adv else 0
                                },
                                'basic_marketing': {
                                    'avg': sum(country_mark)/len(country_mark) if country_mark else 0,
                                    'min': min(country_mark) if country_mark else 0,
                                    'max': max(country_mark) if country_mark else 0,
                                    'median': sorted(country_mark)[len(country_mark)//2] if country_mark else 0
                                },
                                'basic_utility': {
                                    'avg': sum(country_util)/len(country_util) if country_util else 0,
                                    'min': min(country_util) if country_util else 0,
                                    'max': max(country_util) if country_util else 0,
                                    'median': sorted(country_util)[len(country_util)//2] if country_util else 0
                                },
                                'voice_notes_rate': {
                                    'avg': sum(country_voice_notes)/len(country_voice_notes) if country_voice_notes else 0,
                                    'min': min(country_voice_notes) if country_voice_notes else 0,
                                    'max': max(country_voice_notes) if country_voice_notes else 0,
                                    'median': sorted(country_voice_notes)[len(country_voice_notes)//2] if country_voice_notes else 0
                                }
                            },
                            'committed_amount': stat_dict(country_committed),
                            'one_time_dev_cost': stat_dict(dev_costs),
                            'bot_ui_manday_cost': stat_dict(bot_ui_rates),
                            'custom_ai_manday_cost': stat_dict(custom_ai_rates)
                        }
                    # For backward compatibility, keep the old stats[country] as the sum of all regions
                    # Ensure stats[country] has the msg_types structure
                    if 'All' in stats_by_region[country] and 'msg_types' in stats_by_region[country]['All']:
                        stats[country] = stats_by_region[country]['All']
                    else:
                        # Get the first region that has msg_types
                        for region_data in stats_by_region[country].values():
                            if 'msg_types' in region_data:
                                stats[country] = region_data
                                break
                        else:
                            # Fallback: create a basic structure with msg_types
                            stats[country] = {
                                'platform_fee': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                'msg_types': {
                                    'ai': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                    'advanced': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                    'basic_marketing': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                    'basic_utility': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                    'voice_notes_rate': {'avg': 0, 'min': 0, 'max': 0, 'median': 0}
                                },
                                'committed_amount': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                'one_time_dev_cost': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                'bot_ui_manday_cost': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                                'custom_ai_manday_cost': {'avg': 0, 'min': 0, 'max': 0, 'median': 0}
                            }
                # --- Add average discount per country for all message types and manday rates ---
                def avg_discount(chosen_list, rate_card_list):
                    pairs = [
                        (c, r)
                        for c, r in zip(chosen_list, rate_card_list)
                        if isinstance(c, (int, float)) and isinstance(r, (int, float)) and r
                    ]
                    if not pairs:
                        return 0.0
                    return 100 * sum((r - c) / r for c, r in pairs) / len(pairs)
                for country in countries:
                    country_analytics = Analytics.query.filter_by(country=country).all()
                    # Message types
                    ai_chosen = [a.ai_price for a in country_analytics if a.ai_price is not None]
                    ai_rate = [a.ai_rate_card_price for a in country_analytics if a.ai_rate_card_price is not None]
                    adv_chosen = [a.advanced_price for a in country_analytics if a.advanced_price is not None]
                    adv_rate = [a.advanced_rate_card_price for a in country_analytics if a.advanced_rate_card_price is not None]
                    mark_chosen = [a.basic_marketing_price for a in country_analytics if a.basic_marketing_price is not None]
                    mark_rate = [a.basic_marketing_rate_card_price for a in country_analytics if a.basic_marketing_rate_card_price is not None]
                    util_chosen = [a.basic_utility_price for a in country_analytics if a.basic_utility_price is not None]
                    util_rate = [a.basic_utility_rate_card_price for a in country_analytics if a.basic_utility_rate_card_price is not None]
                    # Manday rates
                    bot_ui_chosen = [a.bot_ui_manday_rate for a in country_analytics if a.bot_ui_manday_rate not in (None, 0, '0', '', 'None')]
                    bot_ui_rate = [COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])['bot_ui'] for _ in bot_ui_chosen]
                    custom_ai_chosen = [a.custom_ai_manday_rate for a in country_analytics if a.custom_ai_manday_rate not in (None, 0, '0', '', 'None')]
                    custom_ai_rate = [COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])['custom_ai'] for _ in custom_ai_chosen]
                    stats[country]['avg_discount'] = {
                        'ai': avg_discount(ai_chosen, ai_rate),
                        'advanced': avg_discount(adv_chosen, adv_rate),
                        'basic_marketing': avg_discount(mark_chosen, mark_rate),
                        'basic_utility': avg_discount(util_chosen, util_rate),
                        'bot_ui_manday': avg_discount(bot_ui_chosen, bot_ui_rate),
                        'custom_ai_manday': avg_discount(custom_ai_chosen, custom_ai_rate),
                    }
                # Defensive: Ensure every stat in analytics['stats'] is a dict with msg_types
                for country in list(stats.keys()):
                    if not isinstance(stats[country], dict):
                        stats[country] = {}
                    
                    # Ensure msg_types structure exists
                    if 'msg_types' not in stats[country]:
                        print(f'DEBUG: Adding msg_types fallback for {country}', file=sys.stderr, flush=True)
                        stats[country]['msg_types'] = {
                            'ai': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                            'advanced': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                            'basic_marketing': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                            'basic_utility': {'avg': 0, 'min': 0, 'max': 0, 'median': 0},
                            'voice_notes_rate': {'avg': 0, 'min': 0, 'max': 0, 'median': 0}
                        }
                # Profile-based usage summary (by profile/email)
                profile_usage = []
                profiles = {}
                for a in Analytics.query.all():
                    email = getattr(a, 'user_email', None)
                    name = a.user_name or ''
                    key = email or name or 'Unknown'
                    entry = profiles.setdefault(key, {
                        'email': email,
                        'name': name,
                        'countries': set(),
                        'regions': set(),
                        'count': 0,
                        'last_ts': None,
                    })
                    entry['count'] += 1
                    if a.country:
                        entry['countries'].add(a.country)
                    if a.region:
                        entry['regions'].add(a.region)
                    if a.timestamp and (entry['last_ts'] is None or a.timestamp > entry['last_ts']):
                        entry['last_ts'] = a.timestamp
                for entry in profiles.values():
                    profile_usage.append({
                        'email': entry['email'],
                        'name': entry['name'],
                        'countries': ', '.join(sorted(entry['countries'])),
                        'regions': ', '.join(sorted(entry['regions'])),
                        'calculations': entry['count'],
                        'last_seen': entry['last_ts'],
                    })

                # Account/domain-level usage (group by email domain)
                account_usage = []
                domains = {}
                for a in Analytics.query.all():
                    email = getattr(a, 'user_email', None)
                    if not email or '@' not in email:
                        continue
                    domain = email.split('@', 1)[1].lower()
                    entry = domains.setdefault(domain, {
                        'count': 0,
                        'countries': set(),
                        'regions': set(),
                    })
                    entry['count'] += 1
                    if a.country:
                        entry['countries'].add(a.country)
                    if a.region:
                        entry['regions'].add(a.region)
                for dom, entry in domains.items():
                    account_usage.append({
                        'domain': dom,
                        'calculations': entry['count'],
                        'countries': ', '.join(sorted(entry['countries'])),
                        'regions': ', '.join(sorted(entry['regions'])),
                    })
                # Per-user stats for table
                user_stats = {}
                user_country_currency_list = db.session.query(Analytics.user_name, Analytics.country, Analytics.currency).distinct().all()
                for user, country, currency in user_country_currency_list:
                    user_entries = Analytics.query.filter_by(user_name=user, country=country, currency=currency).all()
                    def get_user_stats(field):
                        vals = [getattr(a, field) for a in user_entries if getattr(a, field) is not None]
                        if vals:
                            return {
                                'avg': sum(vals)/len(vals),
                                'min': min(vals),
                                'max': max(vals),
                                'median': sorted(vals)[len(vals)//2]
                            }
                        else:
                            return {'avg': 0, 'min': 0, 'max': 0, 'median': 0}
                    # Add one_time_dev_cost and per_manday_cost for user/country/currency
                    rates = COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])
                    if country == 'LATAM':
                        if isinstance(rates['bot_ui'], dict):
                            bot_ui_rate = rates['bot_ui'].get(country, list(rates['bot_ui'].values())[0])
                        else:
                            bot_ui_rate = rates['bot_ui']
                        if isinstance(rates['custom_ai'], dict):
                            custom_ai_rate = rates['custom_ai'].get(country, list(rates['custom_ai'].values())[0])
                        else:
                            custom_ai_rate = rates['custom_ai']
                    else:
                        bot_ui_rate = rates['bot_ui']
                        custom_ai_rate = rates['custom_ai']
                    per_manday_cost = (bot_ui_rate + custom_ai_rate) / 2
                    dev_cost = bot_ui_rate + custom_ai_rate
                    one_time_dev_cost_stats = {'avg': dev_cost, 'min': dev_cost, 'max': dev_cost, 'median': dev_cost}
                    per_manday_cost_stats = {'avg': per_manday_cost, 'min': per_manday_cost, 'max': per_manday_cost, 'median': per_manday_cost}
                    user_stats.setdefault(user, {})[(country, currency)] = {
                        'platform_fee': get_user_stats('platform_fee'),
                        'ai': get_user_stats('ai_price'),
                        'advanced': get_user_stats('advanced_price'),
                        'basic_marketing': get_user_stats('basic_marketing_price'),
                        'basic_utility': get_user_stats('basic_utility_price'),
                        'voice_notes_rate': get_user_stats('voice_notes_rate'),
                        'committed_amount': get_user_stats('committed_amount'),
                        'one_time_dev_cost': one_time_dev_cost_stats,
                        'per_manday_cost': per_manday_cost_stats
                    }
                # Compute avg_price_data and avg_platform_fee_data for Chart.js graphs
                avg_price_data = {}
                avg_platform_fee_data = {}
                for country in countries:
                    country_analytics = Analytics.query.filter_by(country=country).all()
                    avg_price_data[country] = {
                        'ai': {
                            'sum': sum(a.ai_price or 0 for a in country_analytics),
                            'count': len([a for a in country_analytics if a.ai_price is not None])
                        },
                        'advanced': {
                            'sum': sum(a.advanced_price or 0 for a in country_analytics),
                            'count': len([a for a in country_analytics if a.advanced_price is not None])
                        },
                        'basic_marketing': {
                            'sum': sum(a.basic_marketing_price or 0 for a in country_analytics),
                            'count': len([a for a in country_analytics if a.basic_marketing_price is not None])
                        },
                        'basic_utility': {
                            'sum': sum(a.basic_utility_price or 0 for a in country_analytics),
                            'count': len([a for a in country_analytics if a.basic_utility_price is not None])
                        }
                    }
                    avg_platform_fee_data[country] = {
                        'sum': sum(a.platform_fee or 0 for a in country_analytics),
                        'count': len([a for a in country_analytics if a.platform_fee is not None])
                    }
                # Top users by number of calculations (show all, not just top 5)
                all_analytics = Analytics.query.all()
                user_names = [row.user_name for row in all_analytics if row.user_name]
                top_users = Counter(user_names).most_common()  # Remove the 5 limit
                # --- AI model & complexity analytics ---
                ai_models = [getattr(a, 'ai_agent_model', None) for a in all_analytics]
                ai_models = [m for m in ai_models if m]
                ai_model_counts = Counter(ai_models).most_common(10)
                ai_complexities = [getattr(a, 'ai_agent_complexity', None) or 'regular' for a in all_analytics]
                ai_complexities = [c for c in ai_complexities if c]
                ai_complexity_counts = Counter(ai_complexities).most_common()
                # Add calculation route counts
                volumes_count = Analytics.query.filter_by(calculation_route='volumes').count()
                bundle_count = Analytics.query.filter_by(calculation_route='bundle').count()
                analytics = {
                    'calculations': total_calculations,
                    'volumes_count': volumes_count,
                    'bundle_count': bundle_count,
                    'calculations_by_day': calculations_by_day,
                    'calculations_by_week': calculations_by_week,
                    'country_counter': country_counter,
                    'platform_fee_stats': {
                        'avg': avg_platform_fee,
                        'min': min_platform_fee,
                        'max': max_platform_fee,
                        'median': median_platform_fee
                    },
                    'platform_fee_options': platform_fee_options,
                    'ai_stats': ai_stats,
                    'advanced_stats': advanced_stats,
                    'marketing_stats': marketing_stats,
                    'utility_stats': utility_stats,
                    'voice_notes_stats': voice_notes_stats,
                    'voice_notes_usage': voice_notes_usage,
                    'voice_notes_percentage': voice_notes_percentage,
                    'message_volumes': message_volumes,
                    'platform_fee_discount_triggered': platform_fee_discount_triggered,
                    'margin_chosen': margin_chosen,
                    'margin_rate_card': margin_rate_card,
                    'stats': stats,
                    'discount_warnings': {},
                    'avg_price_data': avg_price_data,
                    'avg_platform_fee_data': avg_platform_fee_data,
                    'top_users': top_users,
                    'user_stats': user_stats,
                    'all_analytics': all_analytics,
                    'ai_model_counts': ai_model_counts,
                    'ai_complexity_counts': ai_complexity_counts,
                    'stats_by_region': stats_by_region,
                    'profile_usage': profile_usage,
                    'account_usage': account_usage,
                }
                # --- Voice overview analytics ---
                try:
                    voice_q = Analytics.query.filter(Analytics.channel_type.in_(['voice_only','text_voice']))
                    voice_count = voice_q.count()
                    def _agg(col, func_name='avg'):
                        fn = getattr(func, func_name)
                        val = db.session.query(fn(col)).scalar()
                        try:
                            return float(val) if val is not None else 0.0
                        except Exception:
                            return 0.0
                    voice_overview = {
                        'calculations': voice_count,
                        'avg_total_cost': _agg(Analytics.voice_total_cost, 'avg'),
                        'avg_mandays': _agg(Analytics.voice_mandays, 'avg'),
                        'avg_platform_fee': _agg(Analytics.voice_platform_fee, 'avg'),
                        'avg_setup_fee': _agg(Analytics.whatsapp_setup_fee, 'avg'),
                        'sum_wa_outbound_minutes': _agg(Analytics.whatsapp_voice_outbound_minutes, 'sum'),
                        'sum_wa_inbound_minutes': _agg(Analytics.whatsapp_voice_inbound_minutes, 'sum'),
                        'sum_pstn_inbound_minutes': _agg(Analytics.pstn_inbound_ai_minutes, 'sum'),
                        'sum_pstn_outbound_minutes': _agg(Analytics.pstn_outbound_ai_minutes, 'sum'),
                        'sum_pstn_manual_minutes': _agg(Analytics.pstn_manual_minutes, 'sum'),
                        'channel_type_counts': {
                            'voice_only': Analytics.query.filter_by(channel_type='voice_only').count(),
                            'text_voice': Analytics.query.filter_by(channel_type='text_voice').count(),
                        }
                    }
                    analytics['voice_overview'] = voice_overview
                except Exception as _:
                    analytics['voice_overview'] = None
                # --- Advanced Analytics Calculations ---
                # 1. Discount calculations (by type and platform fee)
                def get_discount_stats(price_field, rate_card_field):
                    discounts = []
                    for a in Analytics.query.all():
                        price = getattr(a, price_field, None)
                        rate_card = getattr(a, rate_card_field, None)
                        if price is not None and rate_card is not None and rate_card > 0:
                            discount = (rate_card - price) / rate_card * 100
                            discounts.append(discount)
                    if discounts:
                        avg = sum(discounts) / len(discounts)
                        minv = min(discounts)
                        maxv = max(discounts)
                        median = sorted(discounts)[len(discounts)//2]
                    else:
                        avg = minv = maxv = median = 0
                    # Distribution buckets (0-10%, 10-20%, ...)
                    buckets = [0]*10
                    for d in discounts:
                        idx = min(max(int(d // 10), 0), 9)
                        buckets[idx] += 1
                    return {'avg': avg, 'min': minv, 'max': maxv, 'median': median, 'buckets': buckets, 'count': len(discounts)}

                ai_discount = get_discount_stats('ai_price', 'ai_rate_card_price')
                advanced_discount = get_discount_stats('advanced_price', 'advanced_rate_card_price')
                marketing_discount = get_discount_stats('basic_marketing_price', 'basic_marketing_rate_card_price')
                utility_discount = get_discount_stats('basic_utility_price', 'basic_utility_rate_card_price')
                platform_fee_discount = get_discount_stats('platform_fee', 'platform_fee')  # always 0, but for completeness

                # 2. Most common prices (mode)
                def get_mode(field):
                    vals = [getattr(a, field) for a in Analytics.query.all() if getattr(a, field) is not None]
                    if vals:
                        return Counter(vals).most_common(1)[0][0]
                    return None
                ai_mode = get_mode('ai_price')
                advanced_mode = get_mode('advanced_price')
                marketing_mode = get_mode('basic_marketing_price')
                utility_mode = get_mode('basic_utility_price')
                platform_fee_mode = get_mode('platform_fee')

                # 3. Popular message types (by total volume)
                def get_total_volume(field):
                    return sum(getattr(a, field) or 0 for a in Analytics.query.all())
                total_ai_vol = get_total_volume('ai_volume')
                total_adv_vol = get_total_volume('advanced_volume')
                total_marketing_vol = get_total_volume('basic_marketing_volume')
                total_utility_vol = get_total_volume('basic_utility_volume')
                popular_types = sorted([
                    ('AI', total_ai_vol),
                    ('Advanced', total_adv_vol),
                    ('Marketing', total_marketing_vol),
                    ('Utility', total_utility_vol)
                ], key=lambda x: x[1], reverse=True)

                # 4. Platform fee vs. deal size (correlation)
                import math
                fees = [a.platform_fee for a in Analytics.query.all() if a.platform_fee is not None]
                deal_sizes = [
                    (getattr(a, 'ai_volume') or 0) + (getattr(a, 'advanced_volume') or 0) +
                    (getattr(a, 'basic_marketing_volume') or 0) + (getattr(a, 'basic_utility_volume') or 0)
                    for a in Analytics.query.all()
                ]
                if fees and deal_sizes and len(fees) == len(deal_sizes):
                    n = len(fees)
                    mean_fee = sum(fees)/n
                    mean_deal = sum(deal_sizes)/n
                    cov = sum((fees[i]-mean_fee)*(deal_sizes[i]-mean_deal) for i in range(n)) / n
                    std_fee = math.sqrt(sum((f-mean_fee)**2 for f in fees)/n)
                    std_deal = math.sqrt(sum((d-mean_deal)**2 for d in deal_sizes)/n)
                    correlation = cov / (std_fee*std_deal) if std_fee and std_deal else 0
                else:
                    correlation = 0

                # 5. Seasonality (by month)
                from collections import defaultdict
                month_counts = defaultdict(int)
                for a in Analytics.query.all():
                    if a.timestamp:
                        month = a.timestamp.strftime('%Y-%m')
                        month_counts[month] += 1
                seasonality = dict(sorted(month_counts.items()))

                # Add to analytics dict
                analytics.update({
                    'discounts': {
                        'ai': ai_discount,
                        'advanced': advanced_discount,
                        'marketing': marketing_discount,
                        'utility': utility_discount,
                        'platform_fee': platform_fee_discount
                    },
                    'modes': {
                        'ai': ai_mode,
                        'advanced': advanced_mode,
                        'marketing': marketing_mode,
                        'utility': utility_mode,
                        'platform_fee': platform_fee_mode
                    },
                    'popular_types': popular_types,
                    'platform_fee_vs_deal_correlation': correlation,
                    'seasonality': seasonality
                })
                return render_template('analytics.html', authorized=True, analytics=analytics)
            else:
                flash('Incorrect keyword.', 'error')
                return render_template('analytics.html', authorized=False, analytics={})
        # GET request or any other case
        return render_template('analytics.html', authorized=False, analytics={})
    except Exception as e:
        import traceback
        print("ANALYTICS ERROR:", e, file=sys.stderr, flush=True)
        traceback.print_exc()
        return "Internal Server Error (analytics)", 500

@app.route('/reset-analytics', methods=['POST'])
def reset_analytics():
    """
    Resets the analytics_data dictionary to its initial state.
    """
    db.session.query(Analytics).delete()
    db.session.commit()
    return 'Analytics reset successfully', 200

@app.route('/readme')
def readme():
    """Render the native HTML documentation page."""
    return render_template('readme.html')

# --- Minimal session test routes ---
@app.route('/set-session')
def set_session():
    session['test'] = 'hello'
    return 'Session set!'

@app.route('/get-session')
def get_session():
    return f"Session value: {session.get('test', 'not set')}"

# --- PATCH: Ensure manday rates always included in suggested_prices for 'prices' step ---
def get_default_manday_rates(inputs):
    country = inputs.get('country', 'India') if inputs else 'India'
    dev_location = inputs.get('dev_location', 'India') if inputs else 'India'
    from calculator import COUNTRY_MANDAY_RATES
    rates = COUNTRY_MANDAY_RATES.get(country, COUNTRY_MANDAY_RATES['APAC'])
    if country == 'LATAM':
        default_bot_ui = rates['bot_ui'][dev_location]
        default_custom_ai = rates['custom_ai'][dev_location]
    else:
        default_bot_ui = rates['bot_ui']
        default_custom_ai = rates['custom_ai']
    return default_bot_ui, default_custom_ai

# --- PATCH: Wrap all suggested_prices dicts before rendering 'prices' step ---
def patch_suggested_prices(suggested_prices, inputs):
    default_bot_ui, default_custom_ai = get_default_manday_rates(inputs)
    if 'bot_ui_manday_rate' not in suggested_prices or suggested_prices['bot_ui_manday_rate'] is None:
        suggested_prices['bot_ui_manday_rate'] = default_bot_ui
    if 'custom_ai_manday_rate' not in suggested_prices or suggested_prices['custom_ai_manday_rate'] is None:
        suggested_prices['custom_ai_manday_rate'] = default_custom_ai
    return suggested_prices

@app.route('/analyticsv2', methods=['GET'])
def analytics_v2():
    """
    New advanced analytics and reporting page using analytics.csv.
    """
    return render_template('analyticsv2.html')

def calculate_pricing_simulation(inputs, pricing_inputs=None):
    """
    Returns a dict with detailed calculations for both volume and committed amount routes for the given user inputs.
    """
    try:
        country = inputs.get('country', 'India')
        ai_volume = float(inputs.get('ai_volume', 0) or 0)
        advanced_volume = float(inputs.get('advanced_volume', 0) or 0)
        basic_marketing_volume = float(inputs.get('basic_marketing_volume', 0) or 0)
        basic_utility_volume = float(inputs.get('basic_utility_volume', 0) or 0)
        platform_fee = float(inputs.get('platform_fee', 0) or 0)
        committed_amount = float(inputs.get('committed_amount', 0) or 0)
        from pricing_config import meta_costs_table
        from calculator import get_committed_amount_rate_for_volume
        meta_costs = meta_costs_table.get(country, meta_costs_table['APAC'])
    except Exception as e:
        print(f"ERROR in calculate_pricing_simulation: {e}", file=sys.stderr, flush=True)
        # Return a minimal valid structure to ensure internal section shows
        return {
            'volume_route': {
                'ai_price': 0, 'adv_price': 0, 'mkt_price': 0, 'utl_price': 0,
                'ai_cost': 0, 'adv_cost': 0, 'mkt_cost': 0, 'utl_cost': 0,
                'platform_fee': 0, 'total': 0,
                'ai_volume': 0, 'adv_volume': 0, 'mkt_volume': 0, 'utl_volume': 0,
                'adv_gupshup_rate': 0
            },
            'bundle_route': {
                'ai_price': 0, 'adv_price': 0, 'mkt_price': 0, 'utl_price': 0,
                'committed_amount': 0, 'required_committed_amount': 0, 'nearest_bundle': 0, 'nearest_lower': 0, 'nearest_upper': 0, 'platform_fee': 0, 'total': 0,
                'ai_volume': 0, 'adv_volume': 0, 'mkt_volume': 0, 'utl_volume': 0,
                'ai_included': 0, 'adv_included': 0, 'mkt_included': 0, 'utl_included': 0,
                'ai_lower_included': 0, 'adv_lower_included': 0, 'mkt_lower_included': 0, 'utl_lower_included': 0,
                'ai_upper_included': 0, 'adv_upper_included': 0, 'mkt_upper_included': 0, 'utl_upper_included': 0,
                'ai_required_included': 0, 'adv_required_included': 0, 'mkt_required_included': 0, 'utl_required_included': 0,
                'adv_gupshup_rate': 0, 'lower_tier_rates': None, 'upper_tier_rates': None
            },
            'meta_costs': {'ai': 0, 'advanced': 0, 'marketing': 0, 'utility': 0}
        }
    
    # Use user-chosen prices for both routes
    # First try from inputs, then from pricing_inputs parameter
    ai_price = float(inputs.get('ai_price', 0) or 0)
    adv_price = float(inputs.get('advanced_price', 0) or 0)
    mkt_price = float(inputs.get('basic_marketing_price', 0) or 0)
    utl_price = float(inputs.get('basic_utility_price', 0) or 0)
    
    # If prices are 0 and pricing_inputs is provided, use those
    if pricing_inputs:
        if ai_price == 0:
            ai_price = float(pricing_inputs.get('ai_price', 0) or 0)
        if adv_price == 0:
            adv_price = float(pricing_inputs.get('advanced_price', 0) or 0)
        if mkt_price == 0:
            mkt_price = float(pricing_inputs.get('basic_marketing_price', 0) or 0)
        if utl_price == 0:
            utl_price = float(pricing_inputs.get('basic_utility_price', 0) or 0)
    
    # Calculate Gupshup markup rates for advanced messages
    total_volume = ai_volume + advanced_volume + basic_marketing_volume + basic_utility_volume
    adv_gupshup_rate = get_committed_amount_rate_for_volume(country, 'advanced', total_volume)
    
    ai_final = ai_price + meta_costs['ai']
    adv_final = adv_price + meta_costs['advanced']
    mkt_final = mkt_price + meta_costs['marketing']
    utl_final = utl_price + meta_costs['utility']
    ai_cost = ai_volume * ai_final
    adv_cost = advanced_volume * adv_final
    mkt_cost = basic_marketing_volume * mkt_final
    utl_cost = basic_utility_volume * utl_final
    total = ai_cost + adv_cost + mkt_cost + utl_cost + platform_fee
    # For committed route, use the same user-chosen prices
    ai_final_bundle = ai_final
    adv_final_bundle = adv_final
    mkt_final_bundle = mkt_final
    utl_final_bundle = utl_final
    # Calculate the required committed amount to cover all message volumes
    # This is the minimum committed amount needed to cover the client's actual usage
    # Use user's chosen rates (not bundle rates with meta costs) for fair comparison
    required_committed_amount = (ai_volume * ai_price) + (advanced_volume * adv_price) + (basic_marketing_volume * mkt_price) + (basic_utility_volume * utl_price)
    
    # Find the nearest programmed bundle amount
    from pricing_config import committed_amount_slabs
    slabs = committed_amount_slabs.get(country, committed_amount_slabs['APAC'])
    
    # Get all programmed bundle amounts (both lower and upper bounds of tiers)
    programmed_bundles = []
    for slab in slabs:
        programmed_bundles.append(slab[0])  # Lower bound
        programmed_bundles.append(slab[1])  # Upper bound
    
    # Remove duplicates and sort
    programmed_bundles = sorted(list(set(programmed_bundles)))
    
    # Determine which amount to use for bundle calculations
    # If user chose a committed amount (bundle route), use that; otherwise use required amount
    bundle_calculation_amount = committed_amount if committed_amount > 0 else required_committed_amount
    
    # Find the nearest upper and lower bundles
    # Special case: if bundle_calculation_amount is 0, show meaningful bundle options
    if bundle_calculation_amount == 0:
        # Show the first two meaningful bundle tiers
        nearest_lower = 0  # No bundle
        nearest_upper = programmed_bundles[1] if len(programmed_bundles) > 1 else programmed_bundles[0]  # First meaningful bundle
    else:
        lower_bundles = [bundle for bundle in programmed_bundles if bundle <= bundle_calculation_amount]
        upper_bundles = [bundle for bundle in programmed_bundles if bundle >= bundle_calculation_amount]
        
        nearest_lower = max(lower_bundles) if lower_bundles else programmed_bundles[0]
        nearest_upper = min(upper_bundles) if upper_bundles else programmed_bundles[-1]
    
    # Choose the closer one as the recommended bundle
    distance_to_lower = abs(bundle_calculation_amount - nearest_lower)
    distance_to_upper = abs(bundle_calculation_amount - nearest_upper)
    
    if distance_to_lower <= distance_to_upper:
        nearest_bundle = nearest_lower
    else:
        nearest_bundle = nearest_upper
    
    # Find the tier rates for both bundles
    chosen_tier_rates = None
    lower_tier_rates = None
    upper_tier_rates = None
    
    for slab in slabs:
        if slab[0] <= nearest_bundle < slab[1]:
            chosen_tier_rates = slab[2]
        if slab[0] <= nearest_lower < slab[1]:
            lower_tier_rates = slab[2]
        if slab[0] <= nearest_upper < slab[1]:
            upper_tier_rates = slab[2]
    
    # If committed_amount is 0, use the nearest programmed bundle
    if committed_amount == 0:
        committed_amount = nearest_bundle
        # For committed amount route, use user's chosen rates (not tier rates)
        # The committed amount covers all usage at the user's chosen rates
        ai_final_bundle = ai_price
        adv_final_bundle = adv_price
        mkt_final_bundle = mkt_price
        utl_final_bundle = utl_price
    
    # Calculate how many messages each bundle can actually cover with its amount
    # Using user's chosen rates (not tier rates)
    ai_included = int(nearest_bundle / ai_price) if ai_price > 0 else 0
    adv_included = int(nearest_bundle / adv_price) if adv_price > 0 else 0
    mkt_included = int(nearest_bundle / mkt_price) if mkt_price > 0 else 0
    utl_included = int(nearest_bundle / utl_price) if utl_price > 0 else 0
    
    # Calculate messages covered by lower bundle
    ai_lower_included = int(nearest_lower / ai_price) if ai_price > 0 else 0
    adv_lower_included = int(nearest_lower / adv_price) if adv_price > 0 else 0
    mkt_lower_included = int(nearest_lower / mkt_price) if mkt_price > 0 else 0
    utl_lower_included = int(nearest_lower / utl_price) if utl_price > 0 else 0
    
    # Calculate messages covered by upper bundle
    ai_upper_included = int(nearest_upper / ai_price) if ai_price > 0 else 0
    adv_upper_included = int(nearest_upper / adv_price) if adv_price > 0 else 0
    mkt_upper_included = int(nearest_upper / mkt_price) if mkt_price > 0 else 0
    utl_upper_included = int(nearest_upper / utl_price) if utl_price > 0 else 0
    
    # Calculate messages covered by required amount (use display amount and bundle rates)
    display_amount = committed_amount if committed_amount > 0 else required_committed_amount
    ai_required_included = int(display_amount / ai_final_bundle) if ai_final_bundle > 0 else 0
    adv_required_included = int(display_amount / adv_final_bundle) if adv_final_bundle > 0 else 0
    mkt_required_included = int(display_amount / mkt_final_bundle) if mkt_final_bundle > 0 else 0
    utl_required_included = int(display_amount / utl_final_bundle) if utl_final_bundle > 0 else 0
    
    # No overage calculations for committed amount bundle route
    # The committed amount covers all usage without additional charges
    total_bundle = committed_amount + platform_fee
    
    try:
        return {
            'volume_route': {
                'ai_price': ai_final, 'adv_price': adv_final, 'mkt_price': mkt_final, 'utl_price': utl_final,
                'ai_cost': ai_cost, 'adv_cost': adv_cost, 'mkt_cost': mkt_cost, 'utl_cost': utl_cost,
                'platform_fee': platform_fee, 'total': total,
                'ai_volume': ai_volume, 'adv_volume': advanced_volume, 'mkt_volume': basic_marketing_volume, 'utl_volume': basic_utility_volume,
                'adv_gupshup_rate': adv_gupshup_rate
            },
            'bundle_route': {
                'ai_price': ai_final_bundle, 'adv_price': adv_final_bundle, 'mkt_price': mkt_final_bundle, 'utl_price': utl_final_bundle,
                'committed_amount': committed_amount, 'required_committed_amount': required_committed_amount, 'display_committed_amount': committed_amount if committed_amount > 0 else required_committed_amount, 'nearest_bundle': nearest_bundle, 'nearest_lower': nearest_lower, 'nearest_upper': nearest_upper, 'platform_fee': platform_fee, 'total': total_bundle,
                'ai_volume': ai_volume, 'adv_volume': advanced_volume, 'mkt_volume': basic_marketing_volume, 'utl_volume': basic_utility_volume,
                'ai_included': ai_included, 'adv_included': adv_included, 'mkt_included': mkt_included, 'utl_included': utl_included,
                'ai_lower_included': ai_lower_included, 'adv_lower_included': adv_lower_included, 'mkt_lower_included': mkt_lower_included, 'utl_lower_included': utl_lower_included,
                'ai_upper_included': ai_upper_included, 'adv_upper_included': adv_upper_included, 'mkt_upper_included': mkt_upper_included, 'utl_upper_included': utl_upper_included,
                'ai_required_included': ai_required_included, 'adv_required_included': adv_required_included, 'mkt_required_included': mkt_required_included, 'utl_required_included': utl_required_included,
                'adv_gupshup_rate': adv_gupshup_rate, 'lower_tier_rates': lower_tier_rates, 'upper_tier_rates': upper_tier_rates
            },
            'meta_costs': meta_costs
        }
    except Exception as e:
        print(f"ERROR in calculate_pricing_simulation return: {e}", file=sys.stderr, flush=True)
        # Return a minimal valid structure to ensure internal section shows
        return {
            'volume_route': {
                'ai_price': 0, 'adv_price': 0, 'mkt_price': 0, 'utl_price': 0,
                'ai_cost': 0, 'adv_cost': 0, 'mkt_cost': 0, 'utl_cost': 0,
                'platform_fee': 0, 'total': 0,
                'ai_volume': 0, 'adv_volume': 0, 'mkt_volume': 0, 'utl_volume': 0,
                'adv_gupshup_rate': 0
            },
            'bundle_route': {
                'ai_price': 0, 'adv_price': 0, 'mkt_price': 0, 'utl_price': 0,
                'committed_amount': 0, 'required_committed_amount': 0, 'nearest_bundle': 0, 'nearest_lower': 0, 'nearest_upper': 0, 'platform_fee': 0, 'total': 0,
                'ai_volume': 0, 'adv_volume': 0, 'mkt_volume': 0, 'utl_volume': 0,
                'ai_included': 0, 'adv_included': 0, 'mkt_included': 0, 'utl_included': 0,
                'ai_lower_included': 0, 'adv_lower_included': 0, 'mkt_lower_included': 0, 'utl_lower_included': 0,
                'ai_upper_included': 0, 'adv_upper_included': 0, 'mkt_upper_included': 0, 'utl_upper_included': 0,
                'ai_required_included': 0, 'adv_required_included': 0, 'mkt_required_included': 0, 'utl_required_included': 0,
                'adv_gupshup_rate': 0, 'lower_tier_rates': None, 'upper_tier_rates': None
            },
            'meta_costs': {'ai': 0, 'advanced': 0, 'marketing': 0, 'utility': 0}
        }

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    logger.info(f"Starting Flask app on port {port}")
    print(f"STARTING FLASK APP ON PORT {port}", file=sys.stderr, flush=True)
    app.run(host='0.0.0.0', port=port, debug=True)